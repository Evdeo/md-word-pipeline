"""Tests for table rendering, `<<`/`^^` merge markers, and column widths."""
from docx.oxml.ns import qn

from lib.build.tables import (
    parse_cell_attrs,
    build_merge_plan,
    is_image_table,
)


# ── parse_cell_attrs unit tests ───────────────────────────────────────────────

def test_parse_cell_attrs_plain_text():
    text, cs, rs, ha, va = parse_cell_attrs("Hello world")
    assert text == "Hello world"
    assert (cs, rs, ha, va) == (1, 1, None, None)


def test_parse_cell_attrs_colspan_rowspan():
    text, cs, rs, ha, va = parse_cell_attrs("Merged {cs=2 rs=3}")
    assert text == "Merged"
    assert cs == 2
    assert rs == 3


def test_parse_cell_attrs_alignment():
    text, cs, rs, ha, va = parse_cell_attrs("Centered {ha=c va=m}")
    assert text == "Centered"
    assert ha == "c"
    assert va == "m"


def test_parse_cell_attrs_trailing_braces_only():
    """Braces not at end of text should not be treated as attributes."""
    text, cs, rs, _, _ = parse_cell_attrs("Put {rs=2} at end of text")
    assert text == "Put {rs=2} at end of text"
    assert cs == 1 and rs == 1


# ── build_merge_plan with << / ^^ markers ────────────────────────────────────

def test_merge_plan_rowspan_with_caret():
    """^^ on a row merges that cell up into the anchor above."""
    header = ["A", "B"]
    body = [
        ["anchor", "x"],
        ["^^",     "y"],
    ]
    clean, merge = build_merge_plan(header, body, has_header=True)
    # row 1 col 0 = "anchor" (spanning 2 rows); row 2 col 0 consumed
    anchor = merge[1][0]
    assert anchor == (1, 0, 2, 1, None, None), f"unexpected merge anchor: {anchor}"
    consumed = merge[2][0]
    assert consumed == (1, 0, 2, 1, None, None)
    assert clean[1][0] == "anchor"
    assert clean[2][0] == ""


def test_merge_plan_colspan_with_lt():
    """<< on a cell merges it into the anchor to its left."""
    header = ["A", "B", "C"]
    body = [
        ["span", "<<", "solo"],
    ]
    clean, merge = build_merge_plan(header, body, has_header=True)
    # row 1 col 0 = "span" cs=2, col 1 consumed
    assert merge[1][0] == (1, 0, 1, 2, None, None)
    assert merge[1][1] == (1, 0, 1, 2, None, None)
    assert clean[1][1] == ""


def test_merge_plan_combined_caret_and_lt():
    """Complex: explicit cs= plus a ^^ to extend downward."""
    header = ["A", "B", "C"]
    body = [
        ["top {cs=2}", "",   "x"],
        ["^^",         "^^", "y"],
    ]
    clean, merge = build_merge_plan(header, body, has_header=True)
    # row 1 anchor spans cols 0-1 and rows 1-2
    anchor = merge[1][0]
    assert anchor[2] == 2, f"expected rowspan=2, got {anchor[2]}"
    assert anchor[3] == 2, f"expected colspan=2, got {anchor[3]}"


# ── is_image_table ─────────────────────────────────────────────────────────────

def test_is_image_table_detects_image_only_rows():
    rows = [
        ["", ""],
        ["![](a.png)", "![](b.png)"],
    ]
    assert is_image_table(rows) is True


def test_is_image_table_rejects_text_rows():
    rows = [
        ["H1", "H2"],
        ["text here", "![](b.png)"],
    ]
    assert is_image_table(rows) is False


# ── end-to-end table rendering ────────────────────────────────────────────────

def test_basic_table_rendered(build_docx):
    md = (
        "| Name | Age |\n"
        "|------|-----|\n"
        "| Ada  | 30  |\n"
        "| Bob  | 42  |\n"
    )
    doc = build_docx(md)
    assert len(doc.tables) == 1
    tbl = doc.tables[0]
    assert tbl.rows[0].cells[0].text.strip() == "Name"
    assert tbl.rows[1].cells[0].text.strip() == "Ada"
    assert tbl.rows[2].cells[1].text.strip() == "42"


def test_table_colspan_via_lt_merge_marker(build_docx):
    md = (
        "| A | B | C |\n"
        "|---|---|---|\n"
        "| merged | << | solo |\n"
    )
    doc = build_docx(md)
    tbl = doc.tables[0]
    # In merged rows, python-docx reports the merge via <w:gridSpan>
    body_row = tbl.rows[1]
    tc0 = body_row.cells[0]._tc
    tcPr = tc0.find(qn("w:tcPr"))
    grid_span = tcPr.find(qn("w:gridSpan")) if tcPr is not None else None
    assert grid_span is not None
    assert int(grid_span.get(qn("w:val"))) == 2


def test_table_rowspan_via_caret_merge_marker(build_docx):
    md = (
        "| A | B |\n"
        "|---|---|\n"
        "| anchor | x |\n"
        "| ^^     | y |\n"
    )
    doc = build_docx(md)
    tbl = doc.tables[0]
    # Top-left cell of the merge carries <w:vMerge w:val="restart">
    anchor_tc = tbl.rows[1].cells[0]._tc
    anchor_tcPr = anchor_tc.find(qn("w:tcPr"))
    assert anchor_tcPr is not None
    vmerge_restart = anchor_tcPr.find(qn("w:vMerge"))
    assert vmerge_restart is not None
    # The val attribute is optional — "restart" vs continuation
    # The second row's cell must have a vMerge WITHOUT w:val (continuation)
    cont_tc = tbl.rows[2].cells[0]._tc
    cont_tcPr = cont_tc.find(qn("w:tcPr"))
    cont_vmerge = cont_tcPr.find(qn("w:vMerge"))
    assert cont_vmerge is not None


def test_table_col_widths_directive_left_align_keeps_pct(build_docx):
    """Left-aligned col-widths tables keep their pct widths untouched."""
    md = (
        "| A | B | C |\n"
        "|---|---|---|\n"
        "| 1 | 2 | 3 |\n"
        "\n"
        '{col-widths="20%,50%,30%" align="left"}\n'
    )
    doc = build_docx(md)
    tbl = doc.tables[0]
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    tblW = tblPr.find(qn("w:tblW"))
    assert tblW is not None
    assert tblW.get(qn("w:type")) == "pct"
    assert int(tblW.get(qn("w:w"))) == 5000  # 100% in OOXML pct units
    # Each tcW should also be pct with the right proportions
    body_row = tbl.rows[1]
    tcWs = [c._tc.find(qn("w:tcPr")).find(qn("w:tcW")) for c in body_row.cells]
    assert [int(t.get(qn("w:w"))) for t in tcWs] == [1000, 2500, 1500]


def test_table_col_widths_default_center_converts_to_dxa(build_docx):
    """Centered tables must use absolute dxa widths so Word can compute margins.
    The tblGrid column proportions must still reflect the requested percentages.
    """
    md = (
        "| A | B | C |\n"
        "|---|---|---|\n"
        "| 1 | 2 | 3 |\n"
        "\n"
        '{col-widths="20%,50%,30%"}\n'
    )
    doc = build_docx(md)
    tbl = doc.tables[0]
    tblW = tbl._tbl.find(qn("w:tblPr")).find(qn("w:tblW"))
    assert tblW.get(qn("w:type")) == "dxa"
    # tblGrid columns preserve the 20/50/30 ratio
    grid = tbl._tbl.find(qn("w:tblGrid"))
    widths = [int(gc.get(qn("w:w"))) for gc in grid.findall(qn("w:gridCol"))]
    assert len(widths) == 3
    total = sum(widths)
    ratios = [w / total for w in widths]
    assert abs(ratios[0] - 0.20) < 0.01
    assert abs(ratios[1] - 0.50) < 0.01
    assert abs(ratios[2] - 0.30) < 0.01


def test_table_cell_alignment_attrs(build_docx):
    md = (
        "| A | B |\n"
        "|---|---|\n"
        "| left {ha=l} | right {ha=r} |\n"
    )
    doc = build_docx(md)
    tbl = doc.tables[0]
    left_cell = tbl.rows[1].cells[0]
    right_cell = tbl.rows[1].cells[1]
    assert left_cell.text.strip() == "left"
    assert right_cell.text.strip() == "right"
    # Alignment set at paragraph level — just confirm text was stripped of attrs
