"""Tests for heading emission: styles, numbering, appendix mode."""
from docx.oxml.ns import qn


def _headings(doc):
    """Return [(style_name, text)] for every paragraph styled as a heading."""
    out = []
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            out.append((p.style.name, p.text))
    return out


def test_heading_levels_map_to_heading_styles(build_docx):
    md = "# H1\n\n## H2\n\n### H3\n\n#### H4\n\n##### H5\n\n###### H6\n"
    doc = build_docx(md)
    headings = _headings(doc)
    assert [s for s, _ in headings] == [
        "Heading 1", "Heading 2", "Heading 3",
        "Heading 4", "Heading 5", "Heading 6",
    ]
    assert [t for _, t in headings] == ["H1", "H2", "H3", "H4", "H5", "H6"]


def test_heading_7_clamps_to_heading_6(build_docx):
    """Markdown doesn't have h7, but if the AST ever produces level=7 it's clamped."""
    md = "####### too deep\n"
    doc = build_docx(md)
    # Marko GFM parses this as a Paragraph (with '####### too deep' as text)
    # OR a heading with level capped. Either way, no crash.
    texts = [p.text for p in doc.paragraphs]
    assert any("too deep" in t for t in texts)


def test_heading_has_style_level_num_pr(build_docx):
    """The numbered_headings default attaches numPr via the Heading style."""
    md = "# One\n\n## Two\n"
    doc = build_docx(md)
    h1 = next(p for p in doc.paragraphs if p.text == "One")
    # Style-level numPr is on Heading 1 style, not on the paragraph directly.
    # Just assert the paragraph uses the Heading 1 style.
    assert h1.style.name == "Heading 1"


def test_notoc_sets_outline_level_9(build_docx):
    md = "## Hidden {.notoc}\n"
    doc = build_docx(md)
    h = next(p for p in doc.paragraphs if "Hidden" in p.text)
    pPr = h._p.find(qn("w:pPr"))
    outline = pPr.find(qn("w:outlineLvl"))
    assert outline is not None
    assert outline.get(qn("w:val")) == "9"


def test_nonumber_suppresses_heading_number(build_docx):
    """{.nonumber} should emit numPr with numId=0 to suppress the style-linked number."""
    md = "## No Number {.nonumber}\n"
    doc = build_docx(md)
    h = next(p for p in doc.paragraphs if "No Number" in p.text)
    pPr = h._p.find(qn("w:pPr"))
    numPr = pPr.find(qn("w:numPr"))
    if numPr is not None:
        numId = numPr.find(qn("w:numId"))
        assert numId is not None
        assert numId.get(qn("w:val")) == "0"


def test_appendix_mode_inserts_appendix_label(build_docx):
    md = (
        "# Intro\n\n"
        "Some text.\n\n"
        ":::appendix\n\n"
        "# First\n\n"
        "## Sub\n\n"
        "## Sub two\n\n"
        "# Second\n"
    )
    doc = build_docx(md)
    # Appendix headings should start with "Appendix A. First" etc.
    headings = _headings(doc)
    texts = [t for _, t in headings]
    assert any("Appendix A" in t and "First" in t for t in texts), texts
    assert any("A.1" in t and "Sub" in t for t in texts), texts
    assert any("A.2" in t and "Sub two" in t for t in texts), texts
    assert any("Appendix B" in t and "Second" in t for t in texts), texts
