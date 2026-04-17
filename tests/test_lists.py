"""Tests for list rendering: bullets, ordered, nested, mixed, spacing."""
from docx.oxml.ns import qn


def _list_items(doc):
    """Return [(text, numId, ilvl)] for every paragraph carrying numPr."""
    out = []
    for p in doc.paragraphs:
        pPr = p._p.find(qn("w:pPr"))
        if pPr is None:
            continue
        numPr = pPr.find(qn("w:numPr"))
        if numPr is None:
            continue
        numId_el = numPr.find(qn("w:numId"))
        ilvl_el = numPr.find(qn("w:ilvl"))
        num_id = int(numId_el.get(qn("w:val"))) if numId_el is not None else None
        ilvl = int(ilvl_el.get(qn("w:val"))) if ilvl_el is not None else 0
        out.append((p.text, num_id, ilvl))
    return out


def test_simple_bullet_list(build_docx):
    doc = build_docx("- one\n- two\n- three\n")
    items = _list_items(doc)
    assert [t for t, _, _ in items] == ["one", "two", "three"]
    assert len({nid for _, nid, _ in items}) == 1
    assert all(lvl == 0 for _, _, lvl in items)


def test_simple_ordered_list(build_docx):
    doc = build_docx("1. first\n2. second\n3. third\n")
    items = _list_items(doc)
    assert [t for t, _, _ in items] == ["first", "second", "third"]
    assert len({nid for _, nid, _ in items}) == 1
    assert all(lvl == 0 for _, _, lvl in items)


def test_nested_bullets_depth_three(build_docx):
    md = (
        "- outer\n"
        "    - middle\n"
        "        - inner\n"
    )
    items = _list_items(build_docx(md))
    assert [t for t, _, _ in items] == ["outer", "middle", "inner"]
    assert [lvl for _, _, lvl in items] == [0, 1, 2]
    # Same-type nesting reuses the parent's numId
    assert len({nid for _, nid, _ in items}) == 1


def test_nested_ordered_depth_three(build_docx):
    md = (
        "1. outer\n"
        "    1. middle\n"
        "        1. inner\n"
    )
    items = _list_items(build_docx(md))
    assert [lvl for _, _, lvl in items] == [0, 1, 2]
    assert len({nid for _, nid, _ in items}) == 1


def test_ordered_with_nested_bullets_uses_fresh_numid(build_docx):
    """Bullets nested inside an ordered list must NOT inherit the
    ordered numId — otherwise they render as "1." instead of "•"."""
    md = (
        "1. top\n"
        "    - inner\n"
        "    - inner two\n"
        "2. top two\n"
    )
    items = _list_items(build_docx(md))
    assert [t for t, _, _ in items] == ["top", "inner", "inner two", "top two"]
    top_id, inner_id, inner2_id, top2_id = [nid for _, nid, _ in items]
    assert top_id == top2_id, "outer ordered items share numId"
    assert inner_id == inner2_id, "inner bullets share numId"
    assert inner_id != top_id, "inner bullets get a fresh numId"


def test_bullets_with_nested_ordered_uses_fresh_numid(build_docx):
    md = (
        "- top\n"
        "    1. inner\n"
        "    2. inner two\n"
        "- top two\n"
    )
    items = _list_items(build_docx(md))
    ids = [nid for _, nid, _ in items]
    assert ids[0] == ids[3], "outer bullet items share"
    assert ids[1] == ids[2], "inner ordered items share"
    assert ids[0] != ids[1], "inner ordered gets fresh numId"


def test_list_followed_by_heading_does_not_bleed_numpr(build_docx):
    doc = build_docx("- one\n- two\n\n# Heading\n")
    heading = next((p for p in doc.paragraphs if p.text.endswith("Heading")), None)
    assert heading is not None
    pPr = heading._p.find(qn("w:pPr"))
    if pPr is None:
        return
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        return
    # If a numPr exists it must belong to the heading style (different numId)
    numId_el = numPr.find(qn("w:numId"))
    assert numId_el is not None
    # Grab the list's numId from the preceding paragraph
    list_items = _list_items(doc)
    list_ids = {nid for _, nid, _ in list_items}
    assert int(numId_el.get(qn("w:val"))) not in list_ids, \
        "heading must not reuse the list's numId"


def test_list_item_spacing_is_tight(build_docx):
    doc = build_docx("- a\n- b\n")
    items = _list_items(doc)
    assert len(items) == 2
    for p in doc.paragraphs:
        if p.text not in {"a", "b"}:
            continue
        pPr = p._p.find(qn("w:pPr"))
        sp = pPr.find(qn("w:spacing"))
        assert sp is not None
        assert sp.get(qn("w:before")) == "0"
        assert sp.get(qn("w:after")) == "28"


def test_deep_bullet_nesting_uses_ilvl_not_style_depth(build_docx):
    """Depth 5 — ilvl should climb but the style is capped at level 3."""
    md = (
        "- l0\n"
        "    - l1\n"
        "        - l2\n"
        "            - l3\n"
        "                - l4\n"
    )
    items = _list_items(build_docx(md))
    assert [t for t, _, _ in items] == ["l0", "l1", "l2", "l3", "l4"]
    assert [lvl for _, _, lvl in items] == [0, 1, 2, 3, 4]
