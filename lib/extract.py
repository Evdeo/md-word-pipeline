#!/usr/bin/env python3
"""
docx_to_md.py — Convert an edited Word document back into the markdown source structure.

Output is written to a timestamped folder so existing source files are never overwritten:
  output/imported_YYYY-MM-DD_HH-MM/
    00-frontpage.md        (regenerated shell referencing document-info.yaml)
    document-info.yaml     (existing values + any new revision rows from Word)
    01-introduction.md
    02-next-section.md
    ...
    images/                (extracted images)

Usage:
  python docx_to_md.py                          # Uses output/document.docx + input/
  python docx_to_md.py output/my.docx           # Custom docx
  python docx_to_md.py output/my.docx -i input/ # Custom input dir (for document-info.yaml)
  python docx_to_md.py --force                  # Re-run even if hash unchanged
"""

import argparse
import hashlib
import io
import os
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import yaml

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph as DocxParagraph
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


# ── Constants ─────────────────────────────────────────────────────────────────

EMU_PER_INCH = 914_400

# Heading style name → markdown level
HEADING_STYLE_MAP = {
    "heading 1": 1, "heading 2": 2, "heading 3": 3,
    "heading 4": 4, "heading 5": 5, "heading 6": 6,
}

# Strip leading auto-numbering: "1.", "1.2.", "1.2.3. ", "Appendix A. ", "A1. "
HEADING_NUMBER_RE = re.compile(
    r'^(?:[A-Z]?[\d]+\.)+\s*|^[A-Z]\d*\.\s*'
)

# Strip appendix prefix: "Appendix A" / "A." / "A1." / "A.1 " / "A.1.1 " / "A1.1 "
APPENDIX_PREFIX_RE = re.compile(
    r'^Appendix\s+[A-Z]\b\.?\s*'            # "Appendix A" or "Appendix A. "
    r'|^[A-Z](?:\.\d+|\d+)*\.?\s+'          # "A. " "A.1 " "A.1.1 " "A1 " "A1.1 "
)

# Detect alert label at start of bold run: "[NOTE] ", "[TIP] ", etc.
ALERT_LABEL_RE = re.compile(r'^\[(NOTE|TIP|WARNING|CAUTION|DANGER|SUCCESS)\]\s*$', re.IGNORECASE)


# ── Checksum helpers ──────────────────────────────────────────────────────────

def file_sha256(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def read_stored_hash(hash_path: Path) -> Optional[str]:
    if hash_path.exists():
        return hash_path.read_text(encoding="utf-8").strip()
    return None


def write_hash(hash_path: Path, digest: str):
    hash_path.write_text(digest + "\n", encoding="utf-8")


# ── Filename sanitisation ─────────────────────────────────────────────────────

def heading_to_filename(heading_text: str) -> str:
    """Convert a heading string to a safe lowercase hyphenated filename part."""
    text = heading_text.lower().strip()
    text = text.replace("&", "and")
    text = re.sub(r"[^\w\s-]", "", text)
    text = re.sub(r"[\s_]+", "-", text)
    text = re.sub(r"-+", "-", text).strip("-")
    return text or "section"


def _slug(text: str) -> str:
    """Turn caption text into a short anchor-safe slug."""
    text = text.lower().strip()
    text = re.sub(r"[^\w\s-]", "", text)
    text = re.sub(r"[\s_]+", "-", text)
    return re.sub(r"-+", "-", text).strip("-")[:40]


# ── Image size class mapping ──────────────────────────────────────────────────

def load_size_classes(config_path: Path) -> Dict[str, float]:
    from lib.build.images import build_size_classes, DEFAULT_SIZE_CLASSES
    if config_path.exists():
        with open(config_path, encoding="utf-8") as f:
            cfg = yaml.safe_load(f) or {}
        sizes = build_size_classes(cfg.get("image_sizes"))
    else:
        sizes = dict(DEFAULT_SIZE_CLASSES)
    return sizes


def nearest_size_class(width_emu: int, content_width_emu: int,
                       size_classes: Dict[str, float]) -> str:
    from lib.build.images import nearest_size_class as _nsc
    return _nsc(width_emu, content_width_emu, size_classes)


# ── Reference map (first pass) ────────────────────────────────────────────────

def build_reference_map(doc: Document) -> Dict[str, str]:
    """Scan the entire document body and build bookmark_name → markdown anchor.

    md_to_docx uses deterministic bookmark names:
      fig-<slug>  →  #{slug}
      tbl-<slug>  →  #{slug}
      h-<slug>    →  #{slug}

    We strip the prefix so the markdown anchor matches exactly what
    md_to_docx would have generated from the original source.
    """
    ref_map: Dict[str, str] = {}
    body = doc.element.body
    for elem in body:
        tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        if tag != "p":
            continue
        for bm in elem.iter(qn("w:bookmarkStart")):
            name = bm.get(qn("w:name"), "")
            if not name or name.startswith("_"):
                continue
            if name.startswith("fig-"):
                ref_map[name] = "#" + name[len("fig-"):]
            elif name.startswith("tbl-"):
                ref_map[name] = "#" + name[len("tbl-"):]
            elif name.startswith("h-"):
                ref_map[name] = "#" + name[len("h-"):]
            else:
                ref_map[name] = "#" + name
    return ref_map


# ── Inline run extraction ─────────────────────────────────────────────────────

def _run_text(run) -> str:
    """Convert a single docx Run to markdown inline text."""
    text = run.text
    if not text:
        return ""
    bold   = run.bold
    italic = run.italic
    strike = run.font.strike if run.font else False

    style_name = ""
    try:
        style_name = (run.style.name or "").lower()
    except Exception:
        pass
    is_code = ("courier" in (run.font.name or "").lower() or
               "code" in style_name)

    if is_code:
        return f"`{text}`"
    if bold and italic:
        return f"***{text}***"
    if bold:
        return f"**{text}**"
    if italic:
        return f"*{text}*"
    if strike:
        return f"~~{text}~~"
    return text


def _elem_to_md(elem, ref_map: Dict[str, str]) -> str:
    """Walk a paragraph XML element, resolving hyperlinks via ref_map.

    Plain runs  → formatted markdown text
    w:hyperlink → [display text](#anchor) for internal links
                  [display text](url)     for external links (url from relationship)
    """
    parts = []
    for child in elem:
        child_tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if child_tag == "hyperlink":
            anchor = child.get(qn("w:anchor"), "")
            # Collect all display text inside this hyperlink element
            display = "".join(
                t.text or "" for t in child.iter(qn("w:t"))
            ).strip()

            if anchor:
                md_anchor = ref_map.get(anchor, f"#{anchor}")
                if display:
                    parts.append(f"[{display}]({md_anchor})")
                else:
                    # Fallback: use the anchor slug as link text
                    parts.append(f"[{md_anchor[1:]}]({md_anchor})")
            else:
                # External link — just keep display text (URL not easily available here)
                parts.append(display)

        elif child_tag == "r":
            from docx.text.run import Run
            try:
                run = Run(child, None)
                parts.append(_run_text(run))
            except Exception:
                parts.append("".join(t.text or "" for t in child.iter(qn("w:t"))))

        elif child_tag in ("bookmarkStart", "bookmarkEnd", "proofErr",
                           "rPrChange", "ins", "del"):
            pass  # skip structural / revision-tracking elements

        else:
            parts.append(_elem_to_md(child, ref_map))

    return "".join(parts)


def paragraph_to_md(para: DocxParagraph,
                    ref_map: Optional[Dict[str, str]] = None) -> str:
    """Convert a paragraph to a markdown inline string.

    When ref_map is supplied, internal hyperlinks are reconstructed as
    [display text](#anchor).  Without it only plain run text is returned.
    """
    if ref_map is None:
        return "".join(_run_text(r) for r in para.runs).strip()
    return _elem_to_md(para._p, ref_map).strip()

# ── Caption parsing ───────────────────────────────────────────────────────────

def _parse_caption(para: DocxParagraph) -> Optional[Tuple[str, str, str]]:
    """Parse a Caption-style paragraph.

    Returns (kind, label, description) where:
      kind  = "Figure" | "Table" | "Code Block" | other prefix
      label = e.g. "Figure 1"  (number already in the text)
      description = rest of caption after the colon

    Returns None if parsing fails.
    """
    text = para.text.strip()
    # Caption format built by md_to_docx: "Figure 1: description"
    m = re.match(r'^(Figure|Table|Code Block)\s+(\d+)\s*:\s*(.*)', text, re.IGNORECASE)
    if m:
        kind = m.group(1).title()
        num  = m.group(2)
        desc = m.group(3).strip()
        return kind, f"{kind} {num}", desc
    # Fallback: any "Word: rest" pattern
    m2 = re.match(r'^(\w[\w\s]*?):\s*(.*)', text)
    if m2:
        return m2.group(1).strip(), m2.group(1).strip(), m2.group(2).strip()
    return None


def _caption_to_md(para: DocxParagraph, anchor: Optional[str] = None) -> str:
    """Render a caption paragraph as markdown caption syntax.

    Output: *Figure N: Description text. {#anchor}*
    The {#anchor} is only added when an anchor slug is provided.
    """
    parsed = _parse_caption(para)
    if parsed is None:
        text = para.text.strip()
        return f"*{text}*" if text else ""

    kind, label, desc = parsed
    anchor_part = f" {{#{anchor}}}" if anchor else ""
    return f"*{label}: {desc}{anchor_part}*"


# ── Table extraction ──────────────────────────────────────────────────────────

def _is_borderless_image_table(table: DocxTable) -> bool:
    """True if this looks like a side-by-side image layout table (:::figures)."""
    if len(table.rows) != 1:
        return False
    has_any_image = False
    for cell in table.rows[0].cells:
        has_drawing = cell._tc.find(f'.//{qn("w:drawing")}') is not None
        if has_drawing:
            has_any_image = True
        elif cell.text.strip() and not has_drawing:
            # Non-empty text cell with no image → not an image table
            # (sub-captions a)/b)/c) are text but are allowed alongside images)
            pass
    return has_any_image


def _is_alert_table(table: DocxTable) -> bool:
    """True if this is a 1×1 alert box table produced by md_to_docx."""
    if len(table.rows) != 1 or len(table.rows[0].cells) != 1:
        return False
    cell = table.rows[0].cells[0]
    if not cell.paragraphs:
        return False
    # First run of first paragraph should be a bold [TYPE] label
    first_para = cell.paragraphs[0]
    if not first_para.runs:
        return False
    first_run_text = first_para.runs[0].text.strip()
    return bool(ALERT_LABEL_RE.match(first_run_text))


def _is_revision_table(table: DocxTable) -> bool:
    if not table.rows:
        return False
    header_texts = [c.text.strip().lower() for c in table.rows[0].cells]
    return (len(header_texts) >= 4 and
            "version" in header_texts[0] and
            "date"    in header_texts[1])


def _extract_revision_rows(table: DocxTable) -> List[dict]:
    revisions = []
    for row in table.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) >= 4 and any(cells):
            revisions.append({
                "version": cells[0],
                "date":    cells[1],
                "author":  cells[2],
                "changes": cells[3],
            })
    return revisions


def _alert_to_md(table: DocxTable,
                 ref_map: Optional[Dict[str, str]] = None) -> str:
    """Convert a 1×1 alert table back to GitHub-style alert syntax.

    Output:
      > [!NOTE]
      > Alert body text here.
    """
    cell = table.rows[0].cells[0]
    lines = []
    alert_type = ""

    for para in cell.paragraphs:
        if not para.runs:
            continue
        # Check if first run is the [TYPE] label
        first = para.runs[0].text.strip()
        m = ALERT_LABEL_RE.match(first)
        if m and not alert_type:
            alert_type = m.group(1).upper()
            # Remaining runs on this paragraph are the body text
            body_runs = para.runs[1:]
            body_text = "".join(r.text for r in body_runs).strip()
            if body_text:
                lines.append(body_text)
        else:
            body_text = paragraph_to_md(para, ref_map)
            if body_text:
                lines.append(body_text)

    if not alert_type:
        return ""

    md_lines = [f"> [!{alert_type}]"]
    for line in lines:
        md_lines.append(f"> {line}")
    return "\n".join(md_lines)


def _cell_alignment(cell) -> Optional[str]:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    for para in cell.paragraphs:
        al = para.alignment
        if al == WD_ALIGN_PARAGRAPH.CENTER: return "c"
        if al == WD_ALIGN_PARAGRAPH.RIGHT:  return "r"
        if al == WD_ALIGN_PARAGRAPH.LEFT:   return "l"
    return None


def _cell_valign(cell) -> Optional[str]:
    from docx.enum.table import WD_ALIGN_VERTICAL
    va = cell.vertical_alignment
    if va == WD_ALIGN_VERTICAL.TOP:    return "t"
    if va == WD_ALIGN_VERTICAL.CENTER: return "m"
    if va == WD_ALIGN_VERTICAL.BOTTOM: return "b"
    return None


def table_to_md(table: DocxTable,
                ref_map: Optional[Dict[str, str]] = None) -> str:
    """Convert a docx table to GFM pipe syntax with merge attributes."""
    nrows = len(table.rows)
    if nrows == 0:
        return ""
    ncols = max(len(r.cells) for r in table.rows)

    grid: List[List[Optional[object]]] = []
    for row in table.rows:
        grid_row = list(row.cells)
        while len(grid_row) < ncols:
            grid_row.append(None)
        grid.append(grid_row)

    seen_ids: set = set()
    anchor_map: Dict[Tuple[int,int], Tuple[int,int]] = {}

    for r in range(nrows):
        for c in range(ncols):
            cell = grid[r][c]
            if cell is None:
                continue
            cid = id(cell)
            if cid not in seen_ids:
                seen_ids.add(cid)
                anchor_map[(r, c)] = (r, c)
            else:
                for pr in range(r + 1):
                    for pc in range(ncols):
                        if pr == r and pc >= c:
                            break
                        if grid[pr][pc] is not None and id(grid[pr][pc]) == cid:
                            anchor_map[(r, c)] = anchor_map.get((pr, pc), (pr, pc))
                            break
                    else:
                        continue
                    break

    span_info: Dict[Tuple[int,int], Tuple[int,int]] = {}
    for (r, c), (ar, ac) in anchor_map.items():
        rs, cs = span_info.get((ar, ac), (0, 0))
        span_info[(ar, ac)] = (max(rs, r - ar + 1), max(cs, c - ac + 1))

    md_rows = []
    for r in range(nrows):
        row_cells = []
        for c in range(ncols):
            cell = grid[r][c]
            ar, ac = anchor_map.get((r, c), (r, c))

            if (ar, ac) != (r, c):
                row_cells.append("<<" if ar == r else "^^")
                continue

            if cell is None:
                row_cells.append("")
                continue

            text = paragraph_to_md(cell.paragraphs[0], ref_map) if cell.paragraphs else ""
            rs, cs = span_info.get((r, c), (1, 1))
            ha = _cell_alignment(cell)
            va = _cell_valign(cell)
            attrs = []
            if cs > 1: attrs.append(f"cs={cs}")
            if rs > 1: attrs.append(f"rs={rs}")
            if ha:     attrs.append(f"ha={ha}")
            if va:     attrs.append(f"va={va}")
            if attrs:
                text = f"{text} {{{' '.join(attrs)}}}"
            row_cells.append(text)

        md_rows.append(row_cells)

    if not md_rows:
        return ""

    lines = []
    lines.append("| " + " | ".join(md_rows[0]) + " |")
    lines.append("| " + " | ".join(["---"] * ncols) + " |")
    for row in md_rows[1:]:
        lines.append("| " + " | ".join(row) + " |")

    # Append {col-widths="..."} when columns are non-uniform
    tblGrid = table._tbl.find(f'.//{qn("w:tblGrid")}')
    if tblGrid is not None:
        grid_cols = tblGrid.findall(qn("w:gridCol"))
        raw_widths = [int(c.get(qn("w:w"), 0)) for c in grid_cols]
        total = sum(raw_widths)
        if total > 0 and len(raw_widths) >= 2:
            pcts = [round(w / total * 100) for w in raw_widths]
            # Only output if columns are meaningfully non-uniform
            # (allow ±2% tolerance for rounding)
            min_p, max_p = min(pcts), max(pcts)
            if max_p - min_p > 2:
                col_widths_str = ",".join(f"{p}%" for p in pcts)
                lines.append("")
                lines.append(f'{{col-widths="{col_widths_str}"}}')

    return "\n".join(lines)


# ── Image extraction ──────────────────────────────────────────────────────────

def _image_filename_from_partname(partname: str, counter: int, ext: str) -> str:
    """Derive a meaningful image filename.

    Tries to use the original filename from the relationship partname
    (e.g. /word/media/banner.png → banner.png).  Falls back to a
    sequential name if the partname is generic (image1.png, etc.).
    """
    stem = Path(partname).stem  # e.g. "banner" or "image1"
    # If it's just "imageN" Word auto-assigned, use sequential counter
    if re.match(r'^image\d+$', stem, re.IGNORECASE):
        return f"image{counter:03d}{ext}"
    # Sanitise the original name
    safe = re.sub(r'[^\w.-]', '_', stem)
    return f"{safe}{ext}"


def _extract_images_from_element(element, doc, images_dir: Path,
                                 size_classes: Dict[str, float],
                                 content_width_emu: int,
                                 img_counter: List[int],
                                 hint_name: str = "") -> List[Tuple[str, str, str]]:
    """Find all drawings in element, save images, return [(md_syntax, alt, size_class)].

    hint_name: optional suggested filename stem (e.g. from following caption text).
    """
    from lib.build.overlays import (
        NS_A as _NS_A,
        NS_WPG as _NS_WPG,
        extract_overlay_from_group,
        overlay_to_markdown,
        OverlaySpec,
    )

    results = []
    for drawing in element.findall(f'.//{qn("w:drawing")}'):
        blip = drawing.find(f'.//{qn("a:blip")}')
        if blip is None:
            continue
        r_ns    = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        r_embed = blip.get(f'{{{r_ns}}}embed')
        if not r_embed:
            continue

        # Detect overlay group drawing — graphicData uri == NS_WPG means the
        # drawing is a wpg:wgp group produced by our overlay emitter. The
        # picture inside is the base image; siblings are shapes.
        graphic_data = drawing.find(f'.//{{{_NS_A}}}graphicData')
        is_overlay = (graphic_data is not None and
                      graphic_data.get("uri") == _NS_WPG)

        # Alt text from docPr (descr or title)
        alt = ""
        doc_pr = drawing.find(f'.//{qn("wp:docPr")}')
        if doc_pr is not None:
            alt = doc_pr.get("descr", "") or doc_pr.get("title", "") or ""

        # Rendered width
        extent    = drawing.find(f'.//{qn("wp:extent")}')
        width_emu = int(extent.get("cx", 0)) if extent is not None else 0

        try:
            part      = doc.part.related_parts[r_embed]
            img_bytes = part.blob
            ext       = Path(part.partname).suffix or ".png"
            partname  = str(part.partname)
        except Exception:
            continue

        img_counter[0] += 1

        # Determine filename: prefer hint from caption, then partname, then counter
        if hint_name:
            safe_hint = re.sub(r'[^\w.-]', '_', hint_name)[:40]
            img_name  = f"{safe_hint}{ext}"
        else:
            img_name = _image_filename_from_partname(partname, img_counter[0], ext)

        # Avoid collisions
        img_path = images_dir / img_name
        if img_path.exists():
            stem = img_path.stem
            img_name = f"{stem}_{img_counter[0]}{ext}"
            img_path = images_dir / img_name

        images_dir.mkdir(parents=True, exist_ok=True)
        img_path.write_bytes(img_bytes)

        size_cls = nearest_size_class(width_emu, content_width_emu, size_classes)
        alt_text = alt or img_name

        # Detect paragraph alignment from the parent paragraph element
        # The drawing is inside a w:r inside a w:p — walk up to find jc
        align_cls = ""
        try:
            # element is the paragraph element itself
            pPr = element.find(qn("w:pPr"))
            if pPr is not None:
                jc = pPr.find(qn("w:jc"))
                if jc is not None:
                    jc_val = jc.get(qn("w:val"), "")
                    if jc_val == "center":
                        align_cls = " .center"
                    elif jc_val == "right":
                        align_cls = " .right"
                    # left/start = default, no class needed
        except Exception:
            pass

        if is_overlay:
            # Emit a :::overlay block instead of a plain image
            shapes = extract_overlay_from_group(drawing, width_emu or 1, _height_emu(drawing) or 1)
            spec = OverlaySpec(
                base_src=f"images/{img_name}",
                base_alt=alt_text,
                attrs={"classes": [size_cls]},
                shapes=shapes,
            )
            md = overlay_to_markdown(spec)
        else:
            md = f"![{alt_text}](images/{img_name}){{.{size_cls}{align_cls}}}"
        results.append((md, alt_text, size_cls))

    return results


def _height_emu(drawing) -> int:
    """Read cy from <wp:extent>, falling back to 0."""
    extent = drawing.find(f'.//{qn("wp:extent")}')
    return int(extent.get("cy", 0)) if extent is not None else 0


# ── Paragraph style detection ─────────────────────────────────────────────────

def _style_name(para: DocxParagraph) -> str:
    try:
        return (para.style.name or "").strip()
    except Exception:
        return ""


def _heading_level(para: DocxParagraph) -> Optional[int]:
    sn = _style_name(para).lower()
    return HEADING_STYLE_MAP.get(sn)


def _is_code_para(para: DocxParagraph) -> bool:
    return "code" in _style_name(para).lower()


def _is_caption(para: DocxParagraph) -> bool:
    return "caption" in _style_name(para).lower()


def _is_block_quote(para: DocxParagraph) -> bool:
    return "block quote" in _style_name(para).lower()


def _has_page_break(elem) -> bool:
    """True if the paragraph contains a page break (w:br type=page or lastRenderedPageBreak)."""
    for br in elem.iter(qn("w:br")):
        if br.get(qn("w:type"), "") in ("page", "column"):
            return True
    # Also check sectPr inside pPr → section break acts as page break
    pPr = elem.find(qn("w:pPr"))
    if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
        return True
    return False


def _strip_heading_number(text: str) -> str:
    """Remove auto-generated heading number prefix from heading text.

    Handles:
      "1. Introduction"       → "Introduction"
      "1.2. Background"       → "Background"
      "A. Supported Features" → "Supported Features"   (appendix H1)
      "A.1 GitHub Flavored"   → "GitHub Flavored"      (appendix H2)
      "A.1.1 Table Alignment" → "Table Alignment"      (appendix H3)
      "Appendix A. Title"     → "Title"
    """
    text = HEADING_NUMBER_RE.sub("", text).strip()
    text = APPENDIX_PREFIX_RE.sub("", text).strip()
    return text


def _is_appendix_heading(text: str) -> bool:
    """Return True if the heading text starts with an appendix prefix."""
    return bool(APPENDIX_PREFIX_RE.match(text)) or text.lower().startswith("appendix ")


# ── List handling ─────────────────────────────────────────────────────────────

def _list_info(para: DocxParagraph) -> Optional[Tuple[int, bool]]:
    """Return (ilvl, is_ordered) if this paragraph is a list item, else None.

    Uses both the style name and the numPr ilvl from the XML for accuracy.
    Style name mapping:
      List Bullet / List Bullet 2-5  → unordered (bullet)
      List Number / List Number 2-5  → ordered (numbered)
    """
    from docx.oxml.ns import qn as _qn

    sn = _style_name(para).lower()

    is_bullet  = "list bullet" in sn
    is_number  = "list number" in sn
    if not (is_bullet or is_number):
        return None

    # Prefer explicit ilvl from numPr XML; fall back to style-name digit
    ilvl = 0
    try:
        pPr    = para._p.find(_qn("w:pPr"))
        numPr  = pPr.find(_qn("w:numPr")) if pPr is not None else None
        ilvl_el = numPr.find(_qn("w:ilvl")) if numPr is not None else None
        if ilvl_el is not None:
            ilvl = int(ilvl_el.get(_qn("w:val"), 0))
        else:
            # Fall back: "List Bullet 2" → level 1, "List Bullet 3" → level 2 …
            import re as _re
            m = _re.search(r'\d+$', sn)
            ilvl = (int(m.group()) - 1) if m else 0
    except Exception:
        pass

    return ilvl, is_number


def _collect_list(children: list, start: int, doc,
                  ref_map: Optional[Dict[str, str]] = None) -> Tuple[str, int]:
    """Collect a run of consecutive list paragraphs into indented markdown.

    Returns (markdown_string, last_consumed_index).

    Ordered items use `1.` (GFM auto-numbers), unordered use `-`.
    Nesting is expressed with 2-space indentation per level.
    """
    lines = []
    i = start

    while i < len(children):
        elem = children[i]
        tag  = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        if tag != "p":
            break
        try:
            para = DocxParagraph(elem, doc)
        except Exception:
            break

        info = _list_info(para)
        if info is None:
            break

        ilvl, is_ordered = info
        text = paragraph_to_md(para, ref_map)
        if not text:
            i += 1
            continue

        indent = "  " * ilvl
        marker = "1." if is_ordered else "-"
        lines.append(f"{indent}{marker} {text}")
        i += 1

    return "\n".join(lines), i - 1


# ── Block-quote (non-alert) handling ─────────────────────────────────────────

def _collect_blockquote(children: list, start: int, doc,
                        ref_map: Optional[Dict[str, str]] = None) -> Tuple[str, int]:
    """Collect consecutive Block Quote paragraphs starting at start.

    Returns (markdown_string, last_consumed_index).
    """
    lines = []
    i = start
    while i < len(children):
        elem = children[i]
        tag  = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        if tag != "p":
            break
        try:
            para = DocxParagraph(elem, doc)
        except Exception:
            break
        if not _is_block_quote(para):
            break
        text = paragraph_to_md(para, ref_map)
        if text:
            lines.append(f"> {text}")
        i += 1

    return "\n".join(lines), i - 1


# ── Main document traversal ───────────────────────────────────────────────────

def extract_body_sections(
    doc: Document,
    images_dir: Path,
    size_classes: Dict[str, float],
    content_width_emu: int,
) -> Tuple[List[Tuple[int, str, List[str]]], List[dict]]:
    """Walk document body and extract sections.

    Returns:
        sections  — list of (heading_level, heading_text, [content_lines])
        revisions — list of revision dicts extracted from revision table
    """
    # Build reference map first (single full-document scan)
    ref_map = build_reference_map(doc)

    sections: List[Tuple[int, str, List[str]]] = []
    current_level   = 0
    current_heading = ""
    current_lines:  List[str] = []
    revisions_found: List[dict] = []
    img_counter = [0]
    in_appendix = False

    body     = doc.element.body
    children = list(body)

    def _flush(lines: List[str]) -> List[str]:
        """Ensure there is exactly one trailing blank line."""
        while lines and lines[-1].strip() == "":
            lines.pop()
        lines.append("")
        return lines

    def _append(lines: List[str], text: str):
        """Append a content block ensuring a blank line before it (if non-empty content exists)."""
        if text.strip():
            if lines and lines[-1].strip() != "":
                lines.append("")
            lines.append(text)
            lines.append("")

    i = 0
    while i < len(children):
        elem = children[i]
        tag  = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag

        # ── Top-level sectPr (end of document) ────────────────────────────
        if tag == "sectPr":
            i += 1
            continue

        if tag == "p":
            try:
                para = DocxParagraph(elem, doc)
            except Exception:
                i += 1
                continue

            level    = _heading_level(para)
            raw_text = para.text.strip()

            # ── Page / section break ──────────────────────────────────────
            # Explicit page breaks in non-heading paragraphs → emit marker
            if _has_page_break(elem) and level is None and not raw_text:
                # Blank paragraph that IS a page break
                _append(current_lines, "<!-- pagebreak -->")
                i += 1
                continue

            # ── Images in paragraph ───────────────────────────────────────
            img_results = _extract_images_from_element(
                elem, doc, images_dir, size_classes, content_width_emu, img_counter
            )

            # ── Heading ───────────────────────────────────────────────────
            if level is not None:
                clean = _strip_heading_number(raw_text)

                # Detect entering appendix section
                if level == 1 and raw_text.strip().lower() == "appendix":
                    # "Appendix" H1 is a sentinel — emit :::appendix marker
                    sections.append((current_level, current_heading, _flush(current_lines)))
                    current_level   = 1
                    current_heading = "Appendix"
                    current_lines   = [":::appendix", ""]
                    in_appendix     = True
                    i += 1
                    continue

                if level == 1:
                    # New top-level section
                    sections.append((current_level, current_heading, _flush(current_lines)))
                    current_level   = level
                    current_heading = clean
                    current_lines   = []
                    if _is_appendix_heading(raw_text) and not in_appendix:
                        current_lines = [":::appendix", ""]
                        in_appendix   = True
                else:
                    # Sub-heading (H2-H6)
                    prefix = "#" * level
                    # Check for {.nonumber}: heading has no auto-number prefix and
                    # the raw text equals the clean text (nothing was stripped)
                    was_numbered = (raw_text != clean)
                    attr = "" if was_numbered else " {.nonumber}"
                    _append(current_lines, f"{prefix} {clean}{attr}")

            # ── Caption ───────────────────────────────────────────────────
            elif _is_caption(para) and raw_text:
                cap_md = _caption_to_md(para)
                _append(current_lines, cap_md)

            # ── Code block ────────────────────────────────────────────────
            elif _is_code_para(para) and raw_text:
                code_lines = [para.text]
                while i + 1 < len(children):
                    nxt_elem = children[i + 1]
                    nxt_tag  = nxt_elem.tag.split("}")[-1] if "}" in nxt_elem.tag else nxt_elem.tag
                    if nxt_tag == "p":
                        try:
                            nxt_para = DocxParagraph(nxt_elem, doc)
                            if _is_code_para(nxt_para):
                                code_lines.append(nxt_para.text)
                                i += 1
                                continue
                        except Exception:
                            pass
                    break
                fence  = "```"
                block  = fence + "\n" + "\n".join(code_lines) + "\n" + fence
                _append(current_lines, block)

            # ── List item ─────────────────────────────────────────────────
            elif _list_info(para) is not None:
                list_md, last_i = _collect_list(children, i, doc, ref_map)
                _append(current_lines, list_md)
                i = last_i

            # ── Block quote ───────────────────────────────────────────────
            elif _is_block_quote(para) and raw_text:
                bq_md, last_i = _collect_blockquote(children, i, doc, ref_map)
                _append(current_lines, bq_md)
                i = last_i

            # ── Image paragraph ───────────────────────────────────────────
            elif img_results:
                if len(img_results) == 1:
                    _append(current_lines, img_results[0][0])
                else:
                    fig_lines = [":::figures"]
                    for md, _, _ in img_results:
                        fig_lines.append(md)
                    fig_lines.append(":::")
                    _append(current_lines, "\n".join(fig_lines))

            # ── Regular paragraph ─────────────────────────────────────────
            elif raw_text:
                text = paragraph_to_md(para, ref_map)
                if text:
                    _append(current_lines, text)

            # ── Blank paragraph ───────────────────────────────────────────
            # (ignored — _append already inserts blank lines between blocks)

        elif tag == "tbl":
            try:
                table = DocxTable(elem, doc)
            except Exception:
                i += 1
                continue

            # ── Revision history table ────────────────────────────────────
            if _is_revision_table(table):
                revisions_found = _extract_revision_rows(table)

            # ── Alert box (1×1 table with [TYPE] label) ───────────────────
            elif _is_alert_table(table):
                alert_md = _alert_to_md(table, ref_map)
                if alert_md:
                    _append(current_lines, alert_md)

            # ── Side-by-side image table (:::figures) ─────────────────────
            elif _is_borderless_image_table(table):
                # Check for caption immediately after
                caption_md = ""
                caption_hint = ""
                if i + 1 < len(children):
                    nxt_elem = children[i + 1]
                    nxt_tag  = nxt_elem.tag.split("}")[-1] if "}" in nxt_elem.tag else nxt_elem.tag
                    if nxt_tag == "p":
                        try:
                            nxt_para = DocxParagraph(nxt_elem, doc)
                            if _is_caption(nxt_para):
                                caption_md   = _caption_to_md(nxt_para)
                                caption_hint = _slug(nxt_para.text)
                                i += 1
                        except Exception:
                            pass

                fig_lines = [":::figures"]
                for cell in table.rows[0].cells:
                    img_res = _extract_images_from_element(
                        cell._tc, doc, images_dir, size_classes, content_width_emu,
                        img_counter, hint_name=caption_hint
                    )
                    for md, _, _ in img_res:
                        fig_lines.append(md)
                fig_lines.append(":::")
                _append(current_lines, "\n".join(fig_lines))
                if caption_md:
                    _append(current_lines, caption_md)

            # ── Regular data table ────────────────────────────────────────
            else:
                md_table = table_to_md(table, ref_map)
                if md_table:
                    _append(current_lines, md_table)

        i += 1

    # Flush last section
    sections.append((current_level, current_heading, _flush(current_lines)))

    return sections, revisions_found


# ── document-info.yaml merge ──────────────────────────────────────────────────

def merge_document_info(existing_path: Path, new_revisions: List[dict]) -> dict:
    if existing_path.exists():
        with open(existing_path, encoding="utf-8") as f:
            data = yaml.safe_load(f) or {}
    else:
        data = {"document": {}, "revisions": []}

    existing_versions = {str(r.get("version", "")) for r in data.get("revisions", [])}
    prepend = [r for r in new_revisions if str(r.get("version", "")) not in existing_versions]
    data["revisions"] = prepend + data.get("revisions", [])
    return data


# ── frontpage regeneration ────────────────────────────────────────────────────

FRONTPAGE_TEMPLATE = """\
::: {toc=false align=center size=32 color=#1F3864}
# {{document.title}}
## {{document.subtitle}}
:::

::: {toc=false align=center size=14 color=#666666}
**{{document.document_type}}**
:::

::: {toc=false align=center size=12}
**Version:** {{document.version}} | **Classification:** {{document.classification}}
:::

:::space{lines=6}

{{revisions.table}}
"""


# ── output writer ─────────────────────────────────────────────────────────────

def write_imported(
    out_dir: Path,
    sections: List[Tuple[int, str, List[str]]],
    document_info_data: dict,
):
    out_dir.mkdir(parents=True, exist_ok=True)

    # 00-frontpage.md
    frontpage_path = out_dir / "00-frontpage.md"
    frontpage_path.write_text(FRONTPAGE_TEMPLATE, encoding="utf-8")
    print(f"  Written: {frontpage_path.name}")

    # document-info.yaml
    di_path = out_dir / "document-info.yaml"
    with open(di_path, "w", encoding="utf-8") as f:
        yaml.dump(
            document_info_data, f,
            default_flow_style=False,
            allow_unicode=True,
            sort_keys=False,
        )
    print(f"  Written: {di_path.name}")

    # Content files
    file_index = 1
    for level, heading, lines in sections:
        if level == 0 and not heading:
            continue  # pre-heading content (TOC area etc.)

        safe_name = heading_to_filename(heading)
        filename  = f"{file_index:02d}-{safe_name}.md"
        filepath  = out_dir / filename

        # Build file: H1 heading + blank line + body
        body_text = "\n".join(lines).strip()

        # Remove duplicate leading :::appendix if heading IS "Appendix"
        # (it's encoded in the section heading already)
        if heading.lower() == "appendix" and body_text.startswith(":::appendix"):
            body_text = body_text[len(":::appendix"):].lstrip("\n")

        content = f"# {heading}\n\n{body_text}\n"
        filepath.write_text(content, encoding="utf-8")
        print(f"  Written: {filename}")
        file_index += 1


# ── main ──────────────────────────────────────────────────────────────────────

def main():
    project_root = Path(__file__).parent.resolve()

    ap = argparse.ArgumentParser(
        description="Convert an edited Word document back to markdown source files.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python docx_to_md.py                          # Use defaults
  python docx_to_md.py output/my.docx           # Custom docx
  python docx_to_md.py output/my.docx -i input/ # Custom input dir
  python docx_to_md.py --force                  # Skip hash check
        """
    )
    ap.add_argument("docx", nargs="?", default=None,
                    help="Path to .docx file (default: output/document.docx)")
    ap.add_argument("-i", "--input-dir", default=None,
                    help="Input directory containing document-info.yaml and config.yaml")
    ap.add_argument("--force", action="store_true",
                    help="Run conversion even if the docx hasn't changed")
    args = ap.parse_args()

    docx_path = (Path(args.docx).resolve() if args.docx
                 else (project_root / "output" / "document.docx"))
    input_dir = (Path(args.input_dir).resolve() if args.input_dir
                 else (project_root / "input"))

    if not docx_path.exists():
        print(f"Error: File not found: {docx_path}", file=sys.stderr)
        sys.exit(1)

    hash_path    = docx_path.parent / (docx_path.name + ".hash")
    current_hash = file_sha256(docx_path)
    stored_hash  = read_stored_hash(hash_path)

    if not args.force and stored_hash == current_hash:
        print(f"No changes detected in {docx_path.name} since last conversion.")
        print(f"Use --force to convert anyway.")
        sys.exit(0)

    config_path       = input_dir / "config.yaml"
    size_classes      = load_size_classes(config_path)
    content_width_emu = 170 * 914_400 // 1000  # ~17 cm default

    if config_path.exists():
        with open(config_path, encoding="utf-8") as f:
            cfg = yaml.safe_load(f) or {}
        page_cfg = cfg.get("page", {})
        def cm_to_emu(s):
            try:
                val = float(str(s).replace("cm","").replace("in","").strip())
                if "in" in str(s): return int(val * EMU_PER_INCH)
                return int(val / 2.54 * EMU_PER_INCH)
            except Exception:
                return int(2.54 / 2.54 * EMU_PER_INCH)
        page_size_str = page_cfg.get("size", "A4").upper()
        page_w_emu    = int(21.0  / 2.54 * EMU_PER_INCH) if "A4" in page_size_str else int(21.59 / 2.54 * EMU_PER_INCH)
        ml = cm_to_emu(page_cfg.get("margin_left",  "2.54cm"))
        mr = cm_to_emu(page_cfg.get("margin_right", "2.54cm"))
        content_width_emu = page_w_emu - ml - mr

    timestamp  = datetime.now().strftime("%Y-%m-%d_%H-%M")
    out_dir    = docx_path.parent / f"imported_{timestamp}"
    images_dir = out_dir / "images"

    print(f"\nConverting: {docx_path.name}")
    print(f"Output:     {out_dir}\n")

    doc = Document(str(docx_path))
    sections, revisions_found = extract_body_sections(
        doc, images_dir, size_classes, content_width_emu
    )

    di_path            = input_dir / "document-info.yaml"
    document_info_data = merge_document_info(di_path, revisions_found)

    if revisions_found:
        if di_path.exists():
            existing_versions = {
                str(x.get("version", ""))
                for x in (yaml.safe_load(di_path.read_text()) or {}).get("revisions", [])
            }
        else:
            existing_versions = set()
        new_count = len([r for r in revisions_found
                         if r.get("version") not in existing_versions])
        print(f"Revisions extracted: {len(revisions_found)} rows"
              + (f" ({new_count} new)" if new_count else " (no new entries)"))

    write_imported(out_dir, sections, document_info_data)
    write_hash(hash_path, current_hash)
    print(f"\nHash updated: {hash_path.name}")
    print(f"\nDone. Review files in:\n  {out_dir}")
    print("\nCopy/merge desired changes into your input/ folder.")


if __name__ == "__main__":
    main()
