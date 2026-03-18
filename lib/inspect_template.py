#!/usr/bin/env python3
"""
inspect_template.py — Print style values from any Word (.docx) file.

Run this once against a client or company template, then manually copy
the values you want into convert/styles.py.

Usage:
  python inspect_template.py path/to/template.docx
"""

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


def _rgb_str(color) -> str:
    """Return hex colour string or 'inherited/none'."""
    try:
        if color and color.rgb:
            r, g, b = color.rgb
            return f"#{r:02X}{g:02X}{b:02X}"
    except Exception:
        pass
    return "inherited/none"


def _pt(val) -> str:
    """Return point size as string or 'inherited'."""
    try:
        if val is not None:
            return f"{val.pt:.1f}pt"
    except Exception:
        pass
    return "inherited"


def _bool(val) -> str:
    if val is True:  return "Yes"
    if val is False: return "No"
    return "inherited"


def _indent(n: int) -> str:
    return "  " * n


def _section(title: str):
    print(f"\n{'─' * 60}")
    print(f"  {title}")
    print(f"{'─' * 60}")


def _row(label: str, value: str, indent: int = 1):
    print(f"{_indent(indent)}{label:<28} {value}")


def inspect(docx_path: Path):
    print(f"\n{'═' * 60}")
    print(f"  Template inspection: {docx_path.name}")
    print(f"{'═' * 60}")

    doc = Document(str(docx_path))

    # ── Page layout ───────────────────────────────────────────────────────────
    _section("Page Layout")
    for i, section in enumerate(doc.sections):
        label = f"Section {i + 1}" if len(doc.sections) > 1 else "Default section"
        print(f"\n  {label}")
        pw = section.page_width
        ph = section.page_height
        if pw and ph:
            pw_cm = pw.cm
            ph_cm = ph.cm
            _row("Page size", f"{pw_cm:.2f} cm × {ph_cm:.2f} cm")
            # Identify common sizes
            if abs(pw_cm - 21.0) < 0.3 and abs(ph_cm - 29.7) < 0.3:
                _row("", "→ A4 portrait")
            elif abs(pw_cm - 29.7) < 0.3 and abs(ph_cm - 21.0) < 0.3:
                _row("", "→ A4 landscape")
            elif abs(pw_cm - 21.59) < 0.3 and abs(ph_cm - 27.94) < 0.3:
                _row("", "→ US Letter portrait")

        _row("Margin top",    f"{section.top_margin.cm:.2f} cm" if section.top_margin else "none")
        _row("Margin bottom", f"{section.bottom_margin.cm:.2f} cm" if section.bottom_margin else "none")
        _row("Margin left",   f"{section.left_margin.cm:.2f} cm" if section.left_margin else "none")
        _row("Margin right",  f"{section.right_margin.cm:.2f} cm" if section.right_margin else "none")

    # ── Header / Footer ───────────────────────────────────────────────────────
    _section("Header / Footer")
    section = doc.sections[0]
    for hf_name, hf in [("Header", section.header), ("Footer", section.footer)]:
        texts = [p.text.strip() for p in hf.paragraphs if p.text.strip()]
        if texts:
            print(f"\n  {hf_name}")
            for t in texts:
                _row("Text found", repr(t))
        else:
            print(f"\n  {hf_name}: (empty or linked to previous)")

    # ── Named styles ──────────────────────────────────────────────────────────
    styles_of_interest = [
        "Normal",
        "Heading 1", "Heading 2", "Heading 3",
        "Heading 4", "Heading 5", "Heading 6",
        "Code", "Block Quote", "Caption",
        "Cover Title", "Cover Subtitle", "Cover Body",
    ]

    _section("Paragraph Styles")
    for style_name in styles_of_interest:
        try:
            s = doc.styles[style_name]
        except KeyError:
            print(f"\n  {style_name}: (not present in this template)")
            continue

        print(f"\n  {style_name}")
        f = s.font
        pf = s.paragraph_format

        _row("Font name",     f.name or "inherited")
        _row("Font size",     _pt(f.size))
        _row("Bold",          _bool(f.bold))
        _row("Italic",        _bool(f.italic))
        _row("Color",         _rgb_str(f.color))
        _row("Space before",  _pt(pf.space_before))
        _row("Space after",   _pt(pf.space_after))
        _row("Left indent",   _pt(pf.left_indent))
        try:
            al = pf.alignment
            _row("Alignment", str(al) if al else "inherited")
        except Exception:
            pass

    # ── Table styles ──────────────────────────────────────────────────────────
    _section("Tables (first 3 found in document body)")
    tables_found = 0
    for table in doc.tables:
        if tables_found >= 3:
            break
        tables_found += 1
        print(f"\n  Table {tables_found} ({len(table.rows)} rows × {len(table.columns)} cols)")

        if table.rows:
            hdr_row = table.rows[0]
            print(f"    Header row cells:")
            for ci, cell in enumerate(hdr_row.cells[:4]):
                tc   = cell._tc
                tcPr = tc.find(qn("w:tcPr"))
                fill = "none"
                if tcPr is not None:
                    shd = tcPr.find(qn("w:shd"))
                    if shd is not None:
                        fill = shd.get(qn("w:fill"), "none")
                txt = cell.text.strip()[:30]
                _row(f"  Cell {ci}", f'"{txt}"  bg=#{fill}', indent=2)

        if len(table.rows) > 1:
            body_row = table.rows[1]
            print(f"    First body row cells:")
            for ci, cell in enumerate(body_row.cells[:4]):
                tc   = cell._tc
                tcPr = tc.find(qn("w:tcPr"))
                fill = "none"
                if tcPr is not None:
                    shd = tcPr.find(qn("w:shd"))
                    if shd is not None:
                        fill = shd.get(qn("w:fill"), "none")
                txt = cell.text.strip()[:30]
                _row(f"  Cell {ci}", f'"{txt}"  bg=#{fill}', indent=2)

    if tables_found == 0:
        print("\n  No tables found in document body.")

    # ── Core properties ───────────────────────────────────────────────────────
    _section("Core Document Properties")
    try:
        cp = doc.core_properties
        _row("Title",    cp.title   or "(not set)")
        _row("Author",   cp.author  or "(not set)")
        _row("Subject",  cp.subject or "(not set)")
        _row("Keywords", cp.keywords or "(not set)")
    except Exception as e:
        print(f"  Could not read core properties: {e}")

    print(f"\n{'═' * 60}")
    print("  Done. Copy desired values into convert/styles.py")
    print(f"{'═' * 60}\n")


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    path = Path(sys.argv[1])
    if not path.exists():
        print(f"Error: File not found: {path}", file=sys.stderr)
        sys.exit(1)
    if path.suffix.lower() != ".docx":
        print(f"Warning: Expected a .docx file, got: {path.suffix}")

    inspect(path)


if __name__ == "__main__":
    main()
