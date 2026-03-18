"""
DOCX document builder.

Structure produced:
  Section 1 (cover — 00-frontpage.md):
    - titlePg set  → first-page header/footer are empty
    - pgNumType start = frontpage.cover_start_page (default 1, never visible)
    - Ends with a section-break paragraph

  Section 2 (TOC):
    - Header suppressed when toc_start_page != content_start_page
    - pgNumType start = frontpage.toc_start_page (default 2)
    - Ends with a section-break paragraph after the TOC field

  Section 3 (content):
    - Header always on
    - pgNumType start = frontpage.content_start_page (default 3)
    - {page} counts from content_start_page; {total} = NUMPAGES (whole document)
"""

import re
import unicodedata
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .styles import define_styles
from .images import ImageProcessor, build_size_classes
from .tables import (format_data_table, format_image_table, is_image_table,
                     apply_col_widths, build_merge_plan, apply_merges, parse_cell_attrs)

import marko


# ─── heading numbering XML helpers ────────────────────────────────────────────
# IDs are chosen dynamically at runtime to avoid colliding with the template's
# existing numbering definitions (see _setup_numbering).

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W_PFX = f'xmlns:w="{_W_NS}"'

_LEVEL_FMTS = ["%1.", "%1.%2", "%1.%2.%3",
               "%1.%2.%3.%4", "%1.%2.%3.%4.%5", "%1.%2.%3.%4.%5.%6"]

_HEADING_STYLE_IDS = [
    "Heading1", "Heading2", "Heading3",
    "Heading4", "Heading5", "Heading6",
]


def _abs_num_xml(abs_id: int) -> bytes:
    """Build a w:abstractNum element for multilevel heading numbering.

    Key design points that make this survive Word's on-open "auto-fix":

    • w:nsid / w:tmpl — unique 8-hex identifiers.  Without these Word generates
      its own values and may accidentally match them to another list definition,
      causing counter sharing or style reassignment.

    • w:pStyle in each level — tells Word "level 0 belongs to Heading1, level 1
      belongs to Heading2 …".  This is how Word's own multilevel heading lists
      work.  Without it, Word does not recognise the list as the heading numbering
      and may replace or reset it when the document is opened and fields are
      refreshed.

    • w:suff val="space" — uses a single space after the number instead of a
      variable-width tab, giving consistent number-to-text spacing regardless of
      how many digits the number has.
    """
    lvls = "".join(
        f'<w:lvl w:ilvl="{i}">'
        f'<w:start w:val="1"/><w:numFmt w:val="decimal"/>'
        f'<w:pStyle w:val="{_HEADING_STYLE_IDS[i]}"/>'
        f'<w:lvlText w:val="{fmt}"/><w:lvlJc w:val="left"/>'
        f'<w:suff w:val="space"/>'
        f'<w:pPr><w:ind w:left="0" w:hanging="0"/></w:pPr>'
        f'</w:lvl>'
        for i, fmt in enumerate(_LEVEL_FMTS)
    )
    return (
        f'<w:abstractNum {_W_PFX} w:abstractNumId="{abs_id}">'
        f'<w:nsid w:val="AA00{abs_id:04X}"/>'
        f'<w:multiLevelType w:val="multilevel"/>'
        f'<w:tmpl w:val="BB00{abs_id:04X}"/>'
        f'{lvls}'
        f'</w:abstractNum>'
    ).encode()


def _ol_abs_num_xml(abs_id: int) -> bytes:
    """Build a w:abstractNum element for ordered list numbering.

    Each level is independent (uses only its own counter placeholder) so nested
    lists show 1, 2, 3 — not inherited parent numbers.

    Deliberately NO w:pStyle linkage — this definition is never auto-applied to
    any style.  We attach it explicitly via paragraph-level numPr in _emit_list
    so that each ordered list occurrence gets its own fresh numId (and therefore
    restarts at 1) while remaining completely isolated from heading counters.

    Unique w:nsid / w:tmpl prevent Word from merging this definition with the
    heading or built-in list abstractNums on document open.
    """
    lvls = "".join(
        f'<w:lvl w:ilvl="{i}">'
        f'<w:start w:val="1"/><w:numFmt w:val="decimal"/>'
        f'<w:lvlText w:val="%{i+1}."/><w:lvlJc w:val="left"/>'
        f'<w:suff w:val="space"/>'
        f'<w:pPr><w:ind w:left="{720 + 360*i}" w:hanging="360"/></w:pPr>'
        f'</w:lvl>'
        for i in range(6)
    )
    return (
        f'<w:abstractNum {_W_PFX} w:abstractNumId="{abs_id}">'
        f'<w:nsid w:val="CC00{abs_id:04X}"/>'
        f'<w:multiLevelType w:val="hybridMultilevel"/>'
        f'<w:tmpl w:val="DD00{abs_id:04X}"/>'
        f'{lvls}'
        f'</w:abstractNum>'
    ).encode()


def _ul_abs_num_xml(abs_id: int) -> bytes:
    """Build a w:abstractNum element for unordered (bullet) list numbering.

    Each level uses standard bullet characters with increasing indentation.
    Uses the document's default font (Calibri) instead of Symbol to ensure
    bullets render correctly on all systems.

    Like ordered lists, we use NO w:pStyle linkage and attach numPr explicitly
    in _emit_list so each unordered list gets its own fresh numId and remains
    completely isolated from heading counters.

    Unique w:nsid / w:tmpl prevent Word from merging this definition with the
    heading or other list abstractNums on document open.
    """
    # Standard bullet chars that work in any font: bullet, white circle, small square
    bullet_chars = ['\u2022', '\u25e6', '\u25aa']
    lvls = "".join(
        f'<w:lvl w:ilvl="{i}">'
        f'<w:start w:val="1"/><w:numFmt w:val="bullet"/>'
        f'<w:lvlText w:val="{bullet_chars[i % 3]}"/><w:lvlJc w:val="left"/>'
        f'<w:suff w:val="space"/>'
        f'<w:pPr><w:ind w:left="{720 + 360*i}" w:hanging="360"/></w:pPr>'
        f'</w:lvl>'
        for i in range(6)
    )
    return (
        f'<w:abstractNum {_W_PFX} w:abstractNumId="{abs_id}">'
        f'<w:nsid w:val="EE00{abs_id:04X}"/>'
        f'<w:multiLevelType w:val="hybridMultilevel"/>'
        f'<w:tmpl w:val="FF00{abs_id:04X}"/>'
        f'{lvls}'
        f'</w:abstractNum>'
    ).encode()


def _num_xml(num_id: int, abs_id: int) -> bytes:
    return (
        f'<w:num {_W_PFX} w:numId="{num_id}">'
        f'<w:abstractNumId w:val="{abs_id}"/>'
        f'</w:num>'
    ).encode()


def _num_xml_restart(num_id: int, abs_id: int) -> bytes:
    """Like _num_xml but forces every level's counter to restart at 1.

    Uses w:lvlOverride / w:startOverride — the same mechanism Word itself
    inserts when you click "Restart numbering" in the UI.  This guarantees
    that each new numId begins at 1 regardless of other numIds that share
    the same abstractNum, surviving Word's on-open refresh pass.
    """
    overrides = "".join(
        f'<w:lvlOverride w:ilvl="{i}"><w:startOverride w:val="1"/></w:lvlOverride>'
        for i in range(6)
    )
    return (
        f'<w:num {_W_PFX} w:numId="{num_id}">'
        f'<w:abstractNumId w:val="{abs_id}"/>'
        f'{overrides}'
        f'</w:num>'
    ).encode()


def _full_numbering_xml(heading_abs_id: int, ol_abs_id: int, ul_abs_id: int, num_id: int) -> bytes:
    """Bootstrap numbering.xml from scratch (used when template has no numbering part)."""
    return (
        b"<?xml version='1.0' encoding='UTF-8' standalone='yes'?>"
        b'<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        + _abs_num_xml(heading_abs_id)
        + _ol_abs_num_xml(ol_abs_id)
        + _ul_abs_num_xml(ul_abs_id)
        + _num_xml(num_id, heading_abs_id)
        + b'</w:numbering>'
    )


# ─── helpers ──────────────────────────────────────────────────────────────────

def _slugify(text: str) -> str:
    """Turn heading text into a bookmark-safe identifier."""
    text = unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode()
    text = re.sub(r"[^\w\s-]", "", text).strip().lower()
    return re.sub(r"[\s-]+", "-", text)


def _field_run(para, instr: str):
    """Inject  begin … instrText … end  field code into para."""
    for fldCharType, extra in [("begin", {}), ("separate", {}), ("end", {})]:
        r = OxmlElement("w:r")
        fc = OxmlElement("w:fldChar")
        fc.set(qn("w:fldCharType"), fldCharType)
        if fldCharType == "begin":
            fc.set(qn("w:dirty"), "true")
        r.append(fc)
        para._p.append(r)
        if fldCharType == "begin":
            r2 = OxmlElement("w:r")
            it = OxmlElement("w:instrText")
            it.set(qn("xml:space"), "preserve")
            it.text = instr
            r2.append(it)
            para._p.append(r2)


def _page_num_run(para, instr: str):
    """Add a PAGE or NUMPAGES field as part of a run sequence."""
    r_begin = OxmlElement("w:r")
    fc_b = OxmlElement("w:fldChar"); fc_b.set(qn("w:fldCharType"), "begin")
    r_begin.append(fc_b)

    r_instr = OxmlElement("w:r")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = f" {instr} "
    r_instr.append(it)

    r_sep = OxmlElement("w:r")
    fc_s = OxmlElement("w:fldChar"); fc_s.set(qn("w:fldCharType"), "separate")
    r_sep.append(fc_s)

    r_end = OxmlElement("w:r")
    fc_e = OxmlElement("w:fldChar"); fc_e.set(qn("w:fldCharType"), "end")
    r_end.append(fc_e)

    for r in (r_begin, r_instr, r_sep, r_end):
        para._p.append(r)


def _add_num_pr(para, num_id: int, ilvl: int):
    """Attach list-numbering (numPr) to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    il = OxmlElement("w:ilvl"); il.set(qn("w:val"), str(ilvl))
    ni = OxmlElement("w:numId"); ni.set(qn("w:val"), str(num_id))
    numPr.append(il); numPr.append(ni)
    pPr.insert(0, numPr)


def _suppress_num_pr(para):
    """Explicitly suppress any inherited numbering on this paragraph (numId=0)."""
    pPr = para._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    il = OxmlElement("w:ilvl"); il.set(qn("w:val"), "0")
    ni = OxmlElement("w:numId"); ni.set(qn("w:val"), "0")
    numPr.append(il); numPr.append(ni)
    pPr.insert(0, numPr)


def _add_bookmark(para, bookmark_id: int, name: str):
    """Wrap all content in para with a bookmark."""
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"),   str(bookmark_id))
    bm_start.set(qn("w:name"), name)
    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), str(bookmark_id))
    para._p.insert(0, bm_start)
    para._p.append(bm_end)


# ─── DocumentBuilder ──────────────────────────────────────────────────────────

class DocumentBuilder:

    def __init__(self, config: Optional[Dict] = None, revisions: Optional[list] = None,
                 template_path: Optional[str] = None,
                 source_dir: Optional[Path] = None):
        self.config   = config or {}
        self.revisions = revisions or []
        self._src_dir  = source_dir or Path(".")   # used for header image resolution
        self._fill_defaults()

        self.doc     = Document(template_path) if template_path else Document()
        self._parser = marko.Markdown(extensions=["gfm"])

        page_cfg     = self.config["page"]
        size_classes = build_size_classes(self.config.get("image_sizes"))
        self._img = ImageProcessor(
            page_width_in   = self._page_width_in(page_cfg.get("size", "A4")),
            margin_left_in  = self._cm_to_in(page_cfg.get("margin_left",  "2.54cm")),
            margin_right_in = self._cm_to_in(page_cfg.get("margin_right", "2.54cm")),
            size_classes    = size_classes,
        )

        self._fig_count  = 0          # auto-incrementing figure counter
        self._tbl_count  = 0          # auto-incrementing table counter
        self._bm_id      = 0          # bookmark ID counter
        self._bm_map: Dict[str, str] = {}  # name → bookmark-name
        # Pre-scan label maps: anchor → (type, number)
        # e.g. {"#data-flow": ("Figure", 1), "#feature-status": ("Table", 1)}
        self._label_map: Dict[str, Tuple[str, int]] = {}
        self._last_table          = None                       # for col-widths post-processing
        self._last_img_alignment  = WD_ALIGN_PARAGRAPH.CENTER  # for caption alignment
        self._heading_num_id      = 0     # 0 = disabled; set dynamically once set up
        self._ol_abs_id           = None  # abstractNumId for ordered list numbering
        self._ul_abs_id           = None  # abstractNumId for unordered list numbering
        self._next_num_id         = 0     # next numId to allocate for per-list ordered lists
        self._current_alignment   = None  # alignment override for cover page blocks
        self._current_font_size   = None  # font size override for cover page blocks
        self._current_font_color  = None  # font color override for cover page blocks
        self._last_was_alert      = False  # track if last block was an alert for spacing
        self._appendix_mode       = False  # track if we're in appendix section
        self._appendix_letter     = 'A'    # current appendix letter (A, B, C...)
        self._appendix_counters   = [0, 0, 0, 0, 0, 0]  # counters for H1-H6 in appendix

    # ── label pre-scan ────────────────────────────────────────────────────────

    def prescan_labels(self, all_texts: List[str]):
        """Scan all markdown texts in document order and assign sequential
        numbers to Figure and Table captions that carry a {#anchor} id.

        New caption syntax:
            *Figure: Description text. {#anchor-name}*
            *Table: Description text. {#anchor-name}*

        Populates self._label_map:
            {"#anchor-name": ("Figure", 1), "#other": ("Table", 2), ...}
        """
        fig_num = 0
        tbl_num = 0
        # Pattern: Figure: ... {#name}  or  Table: ... {#name}
        caption_re = re.compile(
            r"^\*?(Figure|Table):\s+.*?\{#([\w-]+)\}\s*\*?$",
            re.MULTILINE
        )
        for text in all_texts:
            for m in caption_re.finditer(text):
                kind   = m.group(1)   # "Figure" or "Table"
                anchor = m.group(2)   # anchor name without #
                if kind == "Figure":
                    fig_num += 1
                    self._label_map[f"#{anchor}"] = ("Figure", fig_num)
                else:
                    tbl_num += 1
                    self._label_map[f"#{anchor}"] = ("Table", tbl_num)

    # ── config helpers ────────────────────────────────────────────────────────

    def _fill_defaults(self):
        # Support both old structure (document: block in config) and new structure
        # where document info comes in via the 'document_info' key (set by md_to_docx.py)
        self.config.setdefault("document", {})
        self.config.setdefault("header",   {"left": "", "center": "", "right": ""})
        self.config.setdefault("footer",   {"left": "", "center": "Page {page} of {total}", "right": ""})
        self.config.setdefault("header_line", {"show": False, "color": "AAAAAA", "width": 6})
        self.config.setdefault("footer_line", {"show": False, "color": "AAAAAA", "width": 6})
        self.config.setdefault("page",     {})
        self.config["page"].setdefault("size",          "A4")
        self.config["page"].setdefault("margin_top",    "2.54cm")
        self.config["page"].setdefault("margin_bottom", "2.54cm")
        self.config["page"].setdefault("margin_left",   "2.54cm")
        self.config["page"].setdefault("margin_right",  "2.54cm")
        self.config.setdefault("numbered_headings", True)
        self.config.setdefault("image_sizes", {})
        fp = self.config.setdefault("frontpage", {})
        fp.setdefault("cover_start_page",   1)   # first page number assigned to cover (hidden)
        fp.setdefault("toc_start_page",     2)   # first page number shown on TOC
        fp.setdefault("content_start_page", 1)   # page number on first content page
        # Note: {total} uses SECTIONPAGES (content section only), so setting
        # content_start_page=1 gives "Page 1 of 20" style numbering throughout.

    @staticmethod
    def _cm_to_in(s: str) -> float:
        if "cm" in s:  return float(s.replace("cm","")) / 2.54
        if "in" in s:  return float(s.replace("in",""))
        return 1.0

    @staticmethod
    def _page_width_in(size: str) -> float:
        return 8.27 if size.upper() == "A4" else 8.5  # A4 vs Letter

    def _in(self, key: str) -> float:
        return self._cm_to_in(self.config["page"].get(key, "2.54cm"))

    # ── public API ────────────────────────────────────────────────────────────

    def setup(self):
        define_styles(self.doc, self.config.get("styles"))
        sec = self.doc.sections[0]
        pc  = self.config["page"]
        sec.top_margin    = Inches(self._cm_to_in(pc.get("margin_top",    "2.54cm")))
        sec.bottom_margin = Inches(self._cm_to_in(pc.get("margin_bottom", "2.54cm")))
        sec.left_margin   = Inches(self._cm_to_in(pc.get("margin_left",   "2.54cm")))
        sec.right_margin  = Inches(self._cm_to_in(pc.get("margin_right",  "2.54cm")))

        # Page orientation — swap width/height for landscape
        orientation = str(pc.get("orientation", "portrait")).lower()
        if orientation == "landscape":
            from docx.oxml.ns import qn as _qn
            from docx.oxml   import OxmlElement as _OE
            pgSz = sec._sectPr.find(_qn("w:pgSz"))
            if pgSz is None:
                pgSz = _OE("w:pgSz")
                sec._sectPr.insert(0, pgSz)
            # Standard A4 landscape: 842pt × 595pt (in twentieths-of-a-point)
            size = str(pc.get("size", "A4")).upper()
            if size == "A4":
                w, h = 16838, 11906   # A4 landscape in twips
            else:
                w, h = 15840, 12240   # Letter landscape
            pgSz.set(_qn("w:w"), str(w))
            pgSz.set(_qn("w:h"), str(h))
            pgSz.set(_qn("w:orient"), "landscape")

        if self.config.get("numbered_headings", True):
            self._setup_numbering()

        # Configure header / footer on the FIRST (and only) section
        self._setup_hdrftr(sec)

    def _setup_numbering(self):
        """Attach multilevel numbering for Heading 1–6, ordered lists, and unordered lists.

        Three abstractNums are created:
          • heading abstractNum       — multilevel %1. / %1.%2 / … used exclusively
                                        by heading paragraphs via explicit numPr.
          • ordered-list abstractNum  — each ordered list gets its OWN numId
                                        pointing here, so each list restarts at 1.
          • unordered-list abstractNum — each unordered (bullet) list gets its OWN numId
                                          pointing here, completely isolated from headings.

        All IDs are chosen dynamically (max + 1) and any numIds referenced by
        styles.xml are also excluded to prevent style-linked counter collisions.
        """
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        from docx.oxml.parser import parse_xml

        NUM_CT  = ("application/vnd.openxmlformats-officedocument"
                   ".wordprocessingml.numbering+xml")
        NUM_REL = ("http://schemas.openxmlformats.org/officeDocument/2006"
                   "/relationships/numbering")

        try:
            # Numbering part already exists — discover max IDs and append
            np   = self.doc.part.numbering_part
            root = np._element
            existing_abs = [
                int(e.get(f'{{{_W_NS}}}abstractNumId', -1))
                for e in root.findall(f'{{{_W_NS}}}abstractNum')
            ]
            existing_num = [
                int(e.get(f'{{{_W_NS}}}numId', -1))
                for e in root.findall(f'{{{_W_NS}}}num')
            ]
            # Also include numIds referenced in styles.xml so we never collide
            # with a style that has a built-in numPr (e.g. List Number → numId=5
            # in python-docx's template).  If our heading numId happened to equal
            # that value, ordered lists would share the heading counter.
            try:
                styles_root = self.doc.part.styles._element
                for e in styles_root.findall(f'.//{{{_W_NS}}}numId'):
                    val = e.get(f'{{{_W_NS}}}val', '-1')
                    if val.isdigit() and int(val) > 0:
                        existing_num.append(int(val))
            except Exception:
                pass

            abs_id    = max(existing_abs, default=-1) + 1   # heading abstractNum
            ol_abs_id = abs_id + 1                           # ordered-list abstractNum
            ul_abs_id = ol_abs_id + 1                        # unordered-list abstractNum
            num_id    = max(existing_num, default=0)  + 1   # heading numId

            root.append(parse_xml(_abs_num_xml(abs_id)))
            root.append(parse_xml(_ol_abs_num_xml(ol_abs_id)))
            root.append(parse_xml(_ul_abs_num_xml(ul_abs_id)))
            root.append(parse_xml(_num_xml(num_id, abs_id)))
            # List numIds are allocated lazily in _new_ol_num_id() and _new_ul_num_id()

        except (AttributeError, KeyError):
            # No numbering part yet — create a fresh one with all three abstractNums
            abs_id, ol_abs_id, ul_abs_id, num_id = 0, 1, 2, 1
            np = Part(
                PackURI("/word/numbering.xml"),
                NUM_CT,
                parse_xml(_full_numbering_xml(abs_id, ol_abs_id, ul_abs_id, num_id)),
                self.doc.part.package,
            )
            self.doc.part.relate_to(np, NUM_REL)

        self._heading_num_id = num_id    # numId for heading paragraphs
        self._ol_abs_id      = ol_abs_id # abstractNumId for ordered-list numIds
        self._ul_abs_id      = ul_abs_id # abstractNumId for unordered-list numIds
        self._next_num_id    = num_id + 1  # start allocating list numIds from here

        # ── Bake numPr into Heading1-6 style definitions ──────────────────────
        # Style-level numPr survives Word's open/refresh cycle reliably.
        # Paragraph-level numPr alone can be silently overwritten by Word when
        # it detects a multilevel list and "auto-fixes" the definition.
        _HEADING_STYLE_NAMES = [
            "Heading 1", "Heading 2", "Heading 3",
            "Heading 4", "Heading 5", "Heading 6",
        ]
        for ilvl, style_id in enumerate(_HEADING_STYLE_IDS):
            try:
                style  = self.doc.styles[_HEADING_STYLE_NAMES[ilvl]]
                s_elem = style._element
                s_pPr  = s_elem.find(qn('w:pPr'))
                if s_pPr is None:
                    s_pPr = OxmlElement('w:pPr')
                    s_elem.append(s_pPr)
                # Remove any existing numPr first
                existing = s_pPr.find(qn('w:numPr'))
                if existing is not None:
                    s_pPr.remove(existing)
                numPr    = OxmlElement('w:numPr')
                ilvl_el  = OxmlElement('w:ilvl');  ilvl_el.set(qn('w:val'),  str(ilvl))
                numid_el = OxmlElement('w:numId'); numid_el.set(qn('w:val'), str(num_id))
                numPr.append(ilvl_el); numPr.append(numid_el)
                s_pPr.insert(0, numPr)
            except Exception:
                pass

    def _new_ol_num_id(self) -> int:
        """Allocate a fresh numId for one ordered-list occurrence.

        Each call creates a new w:num pointing to the ordered-list abstractNum.
        Because each numId has its own independent counter, every ordered list
        automatically restarts at 1 and shares no state with headings or other lists.
        """
        from docx.oxml.parser import parse_xml
        num_id = self._next_num_id
        self._next_num_id += 1
        try:
            np = self.doc.part.numbering_part
            np._element.append(parse_xml(_num_xml_restart(num_id, self._ol_abs_id)))
        except Exception:
            pass
        return num_id

    def _new_ul_num_id(self) -> int:
        """Allocate a fresh numId for one unordered (bullet) list occurrence.

        Each call creates a new w:num pointing to the unordered-list abstractNum.
        Because each numId has its own independent counter, every unordered list
        automatically restarts at 1 and shares no state with headings or other lists.
        """
        from docx.oxml.parser import parse_xml
        num_id = self._next_num_id
        self._next_num_id += 1
        try:
            np = self.doc.part.numbering_part
            np._element.append(parse_xml(_num_xml_restart(num_id, self._ul_abs_id)))
        except Exception:
            pass
        return num_id

    def add_frontpage(self, md_text: str, source_dir: Path):
        self._src_dir = source_dir  # used by header image path resolution
        """Add cover page content then end that page as its own section.
        
        Segments are processed top-to-bottom in document order so that
        :::space{} directives and {{revisions.table}} appear exactly where
        they are written in the source file.
        
        Recognised segment types (each matched at the top level):
          :::  {attrs} … :::   — styled block  (headings, body text)
          :::space{lines=N}    — vertical space (self-closing)
          :::space{pt=N}       — vertical space in points (self-closing)
          {{revisions.table}}  — revision history table
        """
        # Split the text into top-level segments, preserving order.
        # We split on:
        #   1. Styled blocks:  ::: {attrs} ... :::
        #   2. Space blocks:   :::space{...}
        #   3. Revisions:      {{revisions.table}}
        segment_pattern = re.compile(
            r'(:::\s*\{[^}]*\}.*?:::(?:\s*\n)?'   # styled block
            r'|:::space\{[^}]*\}'                  # space directive
            r'|\{\{revisions\.table\}\})',          # revisions placeholder
            re.DOTALL
        )

        for segment in re.split(segment_pattern, md_text):
            segment = segment.strip()
            if not segment:
                continue

            # ── space directive ───────────────────────────────────────────
            m_space = re.match(r'^:::space\{([^}]*)\}$', segment)
            if m_space:
                attrs = m_space.group(1)
                m_lines = re.search(r'lines\s*=\s*(\d+)', attrs)
                m_pt    = re.search(r'pt\s*=\s*([\d.]+)', attrs)
                self._emit_vspace(
                    lines=int(m_lines.group(1)) if m_lines else 0,
                    pt=float(m_pt.group(1)) if m_pt else 0
                )
                continue

            # ── revisions table ───────────────────────────────────────────
            if segment == '{{revisions.table}}':
                self._emit_revisions_table()
                continue

            # ── styled block  ::: {attrs} … ::: ──────────────────────────
            m_block = re.match(r'^:::\s*\{([^}]*)\}(.*?):::$', segment, re.DOTALL)
            if m_block:
                attrs   = m_block.group(1)
                content = m_block.group(2).strip()

                # Alignment
                align_match = re.search(r'align\s*=\s*(\w+)', attrs)
                if align_match:
                    self._current_alignment = {
                        'left':    WD_ALIGN_PARAGRAPH.LEFT,
                        'center':  WD_ALIGN_PARAGRAPH.CENTER,
                        'right':   WD_ALIGN_PARAGRAPH.RIGHT,
                        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
                    }.get(align_match.group(1).lower())
                else:
                    self._current_alignment = None

                # Font size
                size_match = re.search(r'size\s*=\s*(\d+)', attrs)
                self._current_font_size = int(size_match.group(1)) if size_match else None

                # Font color
                color_match = re.search(r'color\s*=\s*#?([0-9A-Fa-f]{6})', attrs)
                if color_match:
                    h = color_match.group(1)
                    self._current_font_color = RGBColor(
                        int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
                    )
                else:
                    self._current_font_color = None

                if content:
                    ast = self._parser.parse(content)
                    for node in ast.children:
                        t = node.__class__.__name__
                        if t == "ThematicBreak":
                            continue
                        elif t == "Heading":
                            self._emit_cover_heading(node, source_dir)
                        else:
                            self._emit_block(node, source_dir)

                self._current_alignment  = None
                self._current_font_size  = None
                self._current_font_color = None

        # Close the cover section in a special paragraph
        # This sectPr ends section 1 (frontpage).
        # titlePg → different first-page header (left empty)
        # pgNumType start=0 → this page counts as 0; TOC = page 1
        sep_para = self.doc.add_paragraph()
        pPr = sep_para._p.get_or_add_pPr()

        sectPr_fp = OxmlElement("w:sectPr")

        titlePg = OxmlElement("w:titlePg")
        sectPr_fp.append(titlePg)

        fp_cfg      = self.config.get("frontpage", {})
        cover_start = int(fp_cfg.get("cover_start_page", 1))
        pgNum = OxmlElement("w:pgNumType")
        pgNum.set(qn("w:fmt"),   "decimal")
        pgNum.set(qn("w:start"), str(cover_start))
        sectPr_fp.append(pgNum)

        # copy margins + set w:header distance so next section inherits correctly
        pc = self.config["page"]
        pgMar = OxmlElement("w:pgMar")
        pgMar.set(qn("w:top"),    str(int(self._cm_to_in(pc.get("margin_top",    "2.54cm")) * 1440)))
        pgMar.set(qn("w:bottom"), str(int(self._cm_to_in(pc.get("margin_bottom", "2.54cm")) * 1440)))
        pgMar.set(qn("w:left"),   str(int(self._cm_to_in(pc.get("margin_left",   "2.54cm")) * 1440)))
        pgMar.set(qn("w:right"),  str(int(self._cm_to_in(pc.get("margin_right",  "2.54cm")) * 1440)))
        # header_distance / footer_distance: distance from paper edge to h/f content
        pc_now = self.config.get("page", {})
        hdr_dist = str(int(self._cm_to_in(pc_now.get("header_distance", "1.25cm")) * 1440))
        ftr_dist = str(int(self._cm_to_in(pc_now.get("footer_distance", "1.25cm")) * 1440))
        pgMar.set(qn("w:header"), hdr_dist)
        pgMar.set(qn("w:footer"), ftr_dist)
        sectPr_fp.append(pgMar)

        pPr.append(sectPr_fp)

    def add_toc(self):
        """Insert TOC field code then close the TOC section.

        Structure after this call:
          [cover section — already closed by add_frontpage]
          TOC heading + TOC field
          [TOC section break paragraph]   ← opens content section
        """
        fp = self.config.get("frontpage", {})
        cover_start   = int(fp.get("cover_start_page",   1))
        toc_start     = int(fp.get("toc_start_page",     2))
        content_start = int(fp.get("content_start_page", 3))
        # Header appears from content_start_page onwards — suppress on TOC if different
        toc_header = (toc_start == content_start)

        heading = self.doc.add_paragraph("Table of Contents", style="Heading 1")
        _suppress_num_pr(heading)
        pPr = heading._p.get_or_add_pPr()
        ol  = OxmlElement("w:outlineLvl"); ol.set(qn("w:val"), "9")
        pPr.append(ol)
        self._bm_id += 1
        _add_bookmark(heading, self._bm_id, "_toc")

        # TOC field
        toc_para = self.doc.add_paragraph()
        _field_run(toc_para, r' TOC \o "1-3" \h \z \u ')

        # ── Close the TOC section, open the content section ───────────────────
        # A paragraph-level sectPr closes the TOC section here.
        # The document-level sectPr (doc.sections[0]) becomes the content section.
        sep = self.doc.add_paragraph()
        pPr = sep._p.get_or_add_pPr()
        sectPr_toc = OxmlElement("w:sectPr")

        # Page numbering on TOC section — starts at toc_start_page
        pgNum_toc = OxmlElement("w:pgNumType")
        pgNum_toc.set(qn("w:fmt"),   "decimal")
        pgNum_toc.set(qn("w:start"), str(toc_start))
        sectPr_toc.append(pgNum_toc)

        # Suppress header on TOC if it starts on a different page number than content
        # (i.e. the user wants the header to only appear from content_start_page)
        if not toc_header:
            titlePg_toc = OxmlElement("w:titlePg")
            sectPr_toc.append(titlePg_toc)

        # Copy margins
        pc = self.config.get("page", {})
        pgMar = OxmlElement("w:pgMar")
        pgMar.set(qn("w:top"),    str(int(self._cm_to_in(pc.get("margin_top",    "2.54cm")) * 1440)))
        pgMar.set(qn("w:bottom"), str(int(self._cm_to_in(pc.get("margin_bottom", "2.54cm")) * 1440)))
        pgMar.set(qn("w:left"),   str(int(self._cm_to_in(pc.get("margin_left",   "2.54cm")) * 1440)))
        pgMar.set(qn("w:right"),  str(int(self._cm_to_in(pc.get("margin_right",  "2.54cm")) * 1440)))
        hdr_dist = str(int(self._cm_to_in(pc.get("header_distance", "1.25cm")) * 1440))
        ftr_dist = str(int(self._cm_to_in(pc.get("footer_distance", "1.25cm")) * 1440))
        pgMar.set(qn("w:header"), hdr_dist)
        pgMar.set(qn("w:footer"), ftr_dist)
        sectPr_toc.append(pgMar)

        pPr.append(sectPr_toc)

        # ── Content section page numbering (on the document-level sectPr) ─────
        # This affects doc.sections[0] which is now the content section.
        # Restart page numbering at content_starts_at.
        # SECTIONPAGES in the footer will count only content-section pages,
        # so {page} of {total} is always self-consistent regardless of cover/TOC length.
        self._content_start = content_start  # stored for _setup_content_section()

    def add_content(self, md_text: str, source_dir: Path):
        self._src_dir = source_dir  # used by header image path resolution
        # Check for :::appendix marker to switch to appendix mode
        if ":::appendix" in md_text:
            # Find the position of :::appendix and split the content
            appendix_pos = md_text.find(":::appendix")
            # Process content before appendix normally
            pre_appendix = md_text[:appendix_pos]
            post_appendix = md_text[appendix_pos + len(":::appendix"):]
            
            # Process pre-appendix content
            self._process_blocks_with_figures(pre_appendix, source_dir)
            
            # Enable appendix mode
            self._appendix_mode = True
            if getattr(self, '_verbose', True):
                print("Entering appendix mode - headings will use Appendix A, A1, etc. numbering")
            
            # Process post-appendix content with appendix numbering
            self._process_blocks_with_figures(post_appendix, source_dir)
        else:
            # No appendix marker - process normally
            self._process_blocks_with_figures(md_text, source_dir)
    
    def _process_blocks_with_figures(self, md_text: str, source_dir: Path):
        """Process markdown blocks, handling :::figures and :::space blocks separately."""
        # Split out :::figures … ::: blocks and :::space{…} blocks before handing text to marko
        # The \n before :::figures ensures we only match blocks at line start,
        # not backtick-quoted instances like `:::figures` in inline text
        # :::space{lines=N} or :::space{pt=N} are self-closing (no ending :::)
        parts = re.split(
            r'(\n:::figures\b[^\n]*\n[\s\S]*?\n:::\n'
            r'|\n:::space\{[^}]*\}\n?)',
            md_text)
        for part in parts:
            if part.startswith("\n:::figures") or part.startswith(":::figures"):
                self._emit_image_group(part.strip(), source_dir)
            elif part.startswith("\n:::space") or part.startswith(":::space"):
                m_lines = re.search(r'lines\s*=\s*(\d+)', part)
                m_pt    = re.search(r'pt\s*=\s*([\d.]+)', part)
                self._emit_vspace(
                    lines=int(m_lines.group(1)) if m_lines else 0,
                    pt=float(m_pt.group(1)) if m_pt else 0
                )
            else:
                self._process_md_blocks(part, source_dir)

    def _process_md_blocks(self, md_text: str, source_dir: Path):
        """Parse and emit normal markdown blocks, with col-widths lookahead."""
        if not md_text.strip():
            return
        ast    = self._parser.parse(md_text)
        blocks = list(ast.children)
        i = 0
        while i < len(blocks):
            node = blocks[i]
            self._emit_block(node, source_dir)

            # If this was a table, look ahead (skipping BlankLines) for a col-widths directive
            # e.g.  {col-widths="20%,50%,30%"}
            if node.__class__.__name__ == "Table":
                j = i + 1
                while j < len(blocks) and blocks[j].__class__.__name__ == "BlankLine":
                    j += 1
                if j < len(blocks) and blocks[j].__class__.__name__ == "Paragraph":
                    nxt_txt = self._plain_text(blocks[j]).strip()
                    m = re.match(r'^\{col-widths="([^"]+)"\}$', nxt_txt)
                    if m and self._last_table is not None:
                        apply_col_widths(self._last_table, m.group(1))
                        i = j + 1
                        continue
            i += 1

        # Ensure the document body never ends with a table.
        # Word always appends an empty paragraph after a table on save —
        # adding one explicitly here prevents that spurious change from
        # showing up as a diff when comparing the built file against a
        # Word-saved version.
        from docx.oxml.ns import qn as _qn
        body = self.doc.element.body
        last_child = None
        for child in body:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag in ('p', 'tbl'):
                last_child = tag
        if last_child == 'tbl':
            self.doc.add_paragraph()


    def add_word_cover(self, cover_path: Path):
        """Insert a Word file as the cover page, then close the cover section.

        All body paragraphs from the cover docx are copied verbatim into the
        main document using lxml XML transplant.  Images and other relationships
        are copied across so they render correctly.  The cover section is then
        closed with the same sectPr used by add_frontpage.
        """
        from docx import Document as _DocX
        from lxml import etree as _et
        import copy as _copy

        cover_doc = _DocX(str(cover_path))

        # ── Copy body paragraphs from cover into main document ────────────────
        body = self.doc.element.body
        cover_body = cover_doc.element.body

        # Copy all paragraphs/tables EXCEPT any trailing sectPr
        for elem in list(cover_body):
            tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
            if tag == "sectPr":
                continue   # skip — we add our own below
            node = _copy.deepcopy(elem)
            # Insert before the document-level sectPr so we don't corrupt structure
            doc_sectPr = body.find(qn("w:sectPr"))
            if doc_sectPr is not None:
                body.insert(list(body).index(doc_sectPr), node)
            else:
                body.append(node)

        # ── Copy image relationships from cover into main document ────────────
        try:
            cover_part = cover_doc.part
            main_part  = self.doc.part
            for rel in cover_part.rels.values():
                if "image" in rel.reltype:
                    try:
                        img_part = rel.target_part
                        # Add image to main document part and record the new rId
                        new_rId = main_part.relate_to(img_part, rel.reltype)
                        # Update rId references in the transplanted XML
                        old_rId = rel.rId
                        if old_rId != new_rId:
                            xml_str = _et.tostring(body, encoding="unicode")
                            xml_str = xml_str.replace(
                                f'r:id="{old_rId}"', f'r:id="{new_rId}"')
                            xml_str = xml_str.replace(
                                f'r:embed="{old_rId}"', f'r:embed="{new_rId}"')
                            new_body = _et.fromstring(xml_str)
                            body.getparent().replace(body, new_body)
                            body = self.doc.element.body
                    except Exception:
                        pass  # best-effort image copy
        except Exception:
            pass  # relationships may not be accessible in all docx variants

        # ── Close the cover section (same logic as add_frontpage) ────────────
        sep_para   = self.doc.add_paragraph()
        pPr        = sep_para._p.get_or_add_pPr()
        sectPr_fp  = OxmlElement("w:sectPr")

        titlePg = OxmlElement("w:titlePg")
        sectPr_fp.append(titlePg)

        fp_cfg      = self.config.get("frontpage", {})
        cover_start = int(fp_cfg.get("cover_start_page", 1))
        pgNum = OxmlElement("w:pgNumType")
        pgNum.set(qn("w:fmt"),   "decimal")
        pgNum.set(qn("w:start"), str(cover_start))
        sectPr_fp.append(pgNum)

        pc = self.config.get("page", {})
        pgMar = OxmlElement("w:pgMar")
        for k, default in [("margin_top", "2.54cm"), ("margin_bottom", "2.54cm"),
                            ("margin_left", "2.54cm"), ("margin_right", "2.54cm")]:
            attr = k.replace("margin_", "")
            pgMar.set(qn(f"w:{attr}"),
                      str(int(self._cm_to_in(pc.get(k, default)) * 1440)))
        hdr_dist = str(int(self._cm_to_in(pc.get("header_distance", "1.25cm")) * 1440))
        ftr_dist = str(int(self._cm_to_in(pc.get("footer_distance", "1.25cm")) * 1440))
        pgMar.set(qn("w:header"), hdr_dist)
        pgMar.set(qn("w:footer"), ftr_dist)
        sectPr_fp.append(pgMar)

        pPr.append(sectPr_fp)

    def save(self, path: Path):
        self._setup_content_section()
        self.doc.save(str(path))

    def _setup_content_section(self):
        """Configure page numbering restart on the content section (doc-level sectPr).

        Called just before save so it runs after all content and section breaks
        have been inserted. The document-level sectPr is now the content section.
        """
        body    = self.doc.element.body
        sectPr  = body.find(qn("w:sectPr"))
        if sectPr is None:
            return

        fp    = self.config.get("frontpage", {})
        start = int(fp.get("content_start_page", getattr(self, "_content_start", 3)))

        # Remove any existing pgNumType and replace with the configured restart
        for old in sectPr.findall(qn("w:pgNumType")):
            sectPr.remove(old)
        pgNum = OxmlElement("w:pgNumType")
        pgNum.set(qn("w:fmt"),   "decimal")
        pgNum.set(qn("w:start"), str(start))
        sectPr.append(pgNum)

    # ── revisions table ──────────────────────────────────────────────────────

    def _emit_revisions_table(self):
        """Render the revisions history as a formatted table on the front page."""
        if not self.revisions:
            return
        
        # Add spacing before table
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(24)
        spacer.paragraph_format.space_after = Pt(12)
        
        # Table title
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("Document Revision History")
        title_run.bold = True
        title_run.font.size = Pt(12)
        title_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        
        # Create table with 4 columns
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Header row
        hdr_cells = table.rows[0].cells
        headers = ["Version", "Date", "Author", "Changes"]
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            # Style header
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            # Set background color
            shading_elm = OxmlElement('w:shd')
            hdr_fill = self.config.get("styles", {}).get("table_header", {}).get("background", "1F3864").lstrip("#")
            shading_elm.set(qn('w:fill'), hdr_fill)
            hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
        
        # Data rows
        for rev in self.revisions:
            row_cells = table.add_row().cells
            row_cells[0].text = rev.get("version", "")
            row_cells[1].text = rev.get("date", "")
            row_cells[2].text = rev.get("author", "")
            row_cells[3].text = rev.get("changes", "")
            
            # Set column widths proportionally
            from docx.shared import Inches
            row_cells[0].width = Inches(0.8)
            row_cells[1].width = Inches(1.2)
            row_cells[2].width = Inches(1.3)
            row_cells[3].width = Inches(3.0)

    # ── vertical space ───────────────────────────────────────────────────────

    def _emit_vspace(self, lines: int = 0, pt: float = 0):
        """Emit vertical space. Lines use body font size (11pt); or exact pt value.
        
        Uses w:lineRule="exact" via direct XML so the paragraph renders at a
        precise fixed height with no visible characters.
        """
        total_pt = pt if pt else lines * 11
        if total_pt <= 0:
            return
        # Word spacing units are twentieths of a point (twips)
        twips = int(total_pt * 20)
        para = self.doc.add_paragraph()
        pPr  = para._p.get_or_add_pPr()
        # Remove any spacing python-docx may have set via style inheritance
        spc = OxmlElement("w:spacing")
        spc.set(qn("w:before"),   "0")
        spc.set(qn("w:after"),    "0")
        spc.set(qn("w:line"),     str(twips))
        spc.set(qn("w:lineRule"), "exact")
        pPr.append(spc)

    # ── header / footer ───────────────────────────────────────────────────────

    def _setup_hdrftr(self, section):
        doc_cfg = self.config["document"]
        title   = doc_cfg.get("title",  "")
        date    = doc_cfg.get("date",   "")
        author  = doc_cfg.get("author", "")

        def subst(t):
            return t.replace("{title}", title).replace("{date}", date).replace("{author}", author)

        def emit_zone(para, text):
            segs = re.split(r"(\{page\}|\{total\})", text)
            for seg in segs:
                if   seg == "{page}":  _page_num_run(para, "PAGE")
                elif seg == "{total}":
                    _ft = self.config.get("footer", {}).get("page_total", "content")
                    _page_num_run(para, "NUMPAGES" if _ft == "document" else "SECTIONPAGES")
                elif seg:              para.add_run(seg)

        def _normalise_zone(val) -> List[str]:
            """Accept either a plain string or a list of strings (multi-line).
            Returns a list of substituted, non-empty strings."""
            if val is None:
                return []
            if isinstance(val, list):
                return [subst(str(v)) for v in val if str(v).strip()]
            s = subst(str(val))
            return [s] if s else []

        def _add_hf_border(para, line_cfg: dict, side: str):
            """Add a top (header) or bottom (footer) paragraph border."""
            color = str(line_cfg.get("color", "AAAAAA")).lstrip("#")
            width = str(line_cfg.get("width", 6))
            pPr  = para._p.get_or_add_pPr()
            existing = pPr.find(qn("w:pBdr"))
            if existing is None:
                existing = OxmlElement("w:pBdr")
                pPr.append(existing)
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"),   "single")
            b.set(qn("w:sz"),    width)
            b.set(qn("w:space"), "4")
            b.set(qn("w:color"), color)
            # Remove existing same-side element first
            for old in existing.findall(qn(f"w:{side}")):
                existing.remove(old)
            existing.append(b)

        def _build_text(hf, cfg, line_cfg: Optional[dict], border_side: str):
            """Build header or footer using tab-based left/center/right layout."""
            left_lines   = _normalise_zone(cfg.get("left",   ""))
            center_lines = _normalise_zone(cfg.get("center", ""))
            right_lines  = _normalise_zone(cfg.get("right",  ""))

            has_content = any([left_lines, center_lines, right_lines])
            show_line   = line_cfg and line_cfg.get("show", False)
            if not has_content and not show_line:
                return

            hf.is_linked_to_previous = False
            n_rows = max(len(left_lines), len(center_lines), len(right_lines), 1)

            def _get_para(idx: int):
                while len(hf.paragraphs) <= idx:
                    hf.add_paragraph()
                p = hf.paragraphs[idx]
                p.clear()
                return p

            def _apply_tabs(para):
                pPr  = para._p.get_or_add_pPr()
                existing = pPr.find(qn("w:tabs"))
                if existing is not None:
                    pPr.remove(existing)
                tabs = OxmlElement("w:tabs")
                for val, pos in [("center", "4680"), ("right", "9360")]:
                    t = OxmlElement("w:tab")
                    t.set(qn("w:val"), val); t.set(qn("w:pos"), pos)
                    tabs.append(t)
                pPr.append(tabs)

            for row in range(n_rows):
                para = _get_para(row)
                l = left_lines[row]   if row < len(left_lines)   else ""
                c = center_lines[row] if row < len(center_lines) else ""
                r = right_lines[row]  if row < len(right_lines)  else ""
                emit_zone(para, l)
                para.add_run("\t")
                emit_zone(para, c)
                para.add_run("\t")
                emit_zone(para, r)
                _apply_tabs(para)

            if show_line:
                _add_hf_border(hf.paragraphs[n_rows - 1], line_cfg, border_side)

        def _build_header_with_image(hf, cfg, line_cfg: Optional[dict],
                                      img_path: Path):
            """Build header as a two-column borderless table.

            Left column  : text rows (left / center content, multi-line)
            Right column : image, vertically centred, right-aligned

            The separator line is applied to the last paragraph of the
            left cell so it only spans the text area, not the image cell.
            """
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.enum.text  import WD_ALIGN_PARAGRAPH
            from docx.oxml      import OxmlElement
            from docx.shared    import Pt

            hf.is_linked_to_previous = False

            # Clear default paragraph that Word puts in every header
            for p in list(hf.paragraphs):
                p._element.getparent().remove(p._element)

            # ── Measure image ─────────────────────────────────────────────────
            # img_path is already fully resolved — pass it as absolute so
            # ImageProcessor.load() doesn't try to join it with base_dir again
            result = self._img.load(str(img_path.resolve()), img_path.parent)
            if result is None:
                # Image not found — fall back to text-only header
                _build_text(hf, cfg, line_cfg, "bottom")
                return
            _, w_px, h_px = result

            # Size image by HEIGHT so it never overflows the header row.
            # header.image_height_cm controls this; default 1.0cm fits a
            # standard two-line header comfortably.
            height_cm  = float(str(cfg.get("image_height_cm", "1.0")).replace("cm",""))
            h_emu      = int(height_cm / 2.54 * 914_400)
            # Scale width proportionally from the original pixel dimensions
            aspect     = w_px / h_px if h_px else 1.0
            w_emu      = int(h_emu * aspect)
            img_in     = self._img.inches(w_emu)

            # ── Page and content dimensions ───────────────────────────────────
            content_w_emu = self._img.content_width_emu
            content_w_in  = self._img.inches(content_w_emu)
            text_w_in     = max(content_w_in - img_in - 0.15, content_w_in * 0.55)

            # ── Build table ───────────────────────────────────────────────────
            tbl = hf.add_table(rows=1, cols=2, width=Inches(content_w_in))
            # "Table Normal" has no default borders/spacing, cleaner than Table Grid
            # No style — borders cleared manually below

            def _make_none_border():
                b = OxmlElement("w:tcBdr")
                for side in ("top","left","bottom","right","insideH","insideV"):
                    el = OxmlElement(f"w:{side}")
                    el.set(qn("w:val"), "none")
                    el.set(qn("w:sz"), "0")
                    el.set(qn("w:space"), "0")
                    el.set(qn("w:color"), "auto")
                    b.append(el)
                return b

            def _no_borders(cell):
                tc   = cell._tc
                tcPr = tc.get_or_add_tcPr()
                for old in tcPr.findall(qn("w:tcBdr")):
                    tcPr.remove(old)
                tcPr.append(_make_none_border())

            # Remove table-level borders via tblPr XML element
            tbl_elem = tbl._tbl
            tblPr_el = tbl_elem.tblPr          # CT_Tbl.tblPr — always exists
            tblBdr   = OxmlElement("w:tblBdr")
            for side in ("top","left","bottom","right","insideH","insideV"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "none"); b.set(qn("w:sz"), "0")
                b.set(qn("w:space"), "0"); b.set(qn("w:color"), "auto")
                tblBdr.append(b)
            for old in tblPr_el.findall(qn("w:tblBdr")):
                tblPr_el.remove(old)
            tblPr_el.append(tblBdr)

            # Zero table indent so it aligns flush with the page margins
            for old in tblPr_el.findall(qn("w:tblInd")):
                tblPr_el.remove(old)
            ind = OxmlElement("w:tblInd")
            ind.set(qn("w:w"), "0"); ind.set(qn("w:type"), "dxa")
            tblPr_el.append(ind)

            # Set column widths (twips = 1440 per inch)
            def _set_col_width(cell, width_in: float):
                tc   = cell._tc
                tcPr = tc.get_or_add_tcPr()
                for old in tcPr.findall(qn("w:tcW")):
                    tcPr.remove(old)
                tcW = OxmlElement("w:tcW")
                tcW.set(qn("w:w"), str(int(width_in * 1440)))
                tcW.set(qn("w:type"), "dxa")
                tcPr.append(tcW)

            # image_position: "left" puts the image in col 0, text in col 1.
            # Anything else (default "right") puts text in col 0, image in col 1.
            img_pos   = str(cfg.get("image_position", "right")).lower()
            img_col   = 0 if img_pos == "left" else 1
            text_col  = 1 if img_pos == "left" else 0

            img_cell  = tbl.cell(0, img_col)
            text_cell = tbl.cell(0, text_col)
            _no_borders(text_cell)
            _no_borders(img_cell)
            # Text cell gets the wider allocation; image cell gets exact image width
            if img_pos == "left":
                _set_col_width(img_cell,  img_in + 0.05)
                _set_col_width(text_cell, text_w_in)
            else:
                _set_col_width(text_cell, text_w_in)
                _set_col_width(img_cell,  img_in + 0.05)

            # ── Text cell — fill in rows ──────────────────────────────────────
            left_lines   = _normalise_zone(cfg.get("left",   ""))
            center_lines = _normalise_zone(cfg.get("center", ""))
            right_lines  = _normalise_zone(cfg.get("right",  ""))
            n_rows = max(len(left_lines), len(center_lines), len(right_lines), 1)

            def _cell_tab_stops(para, col_w_in: float):
                """Set center and right tab stops relative to the cell width."""
                pPr  = para._p.get_or_add_pPr()
                for old in pPr.findall(qn("w:tabs")):
                    pPr.remove(old)
                tabs = OxmlElement("w:tabs")
                for val, frac in [("center", 0.5), ("right", 1.0)]:
                    t = OxmlElement("w:tab")
                    t.set(qn("w:val"), val)
                    t.set(qn("w:pos"), str(int(col_w_in * frac * 1440)))
                    tabs.append(t)
                pPr.append(tabs)

            # Clear the default paragraph Word adds to every new cell
            for p in list(text_cell.paragraphs):
                p._element.getparent().remove(p._element)

            for row in range(n_rows):
                para = text_cell.add_paragraph()
                l = left_lines[row]   if row < len(left_lines)   else ""
                c = center_lines[row] if row < len(center_lines) else ""
                r = right_lines[row]  if row < len(right_lines)  else ""
                emit_zone(para, l)
                if c:
                    para.add_run("\t")
                    emit_zone(para, c)
                if r:
                    para.add_run("\t")
                    emit_zone(para, r)
                _cell_tab_stops(para, text_w_in)
                # Reduce paragraph spacing inside header table
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after  = Pt(0)

            show_line = line_cfg and line_cfg.get("show", False)
            if show_line and text_cell.paragraphs:
                _add_hf_border(text_cell.paragraphs[-1], line_cfg, "bottom")

            # ── Image cell — right-aligned image ─────────────────────────────
            for p in list(img_cell.paragraphs):
                p._element.getparent().remove(p._element)

            img_para = img_cell.add_paragraph()
            img_para.alignment = (WD_ALIGN_PARAGRAPH.LEFT
                                  if img_pos == "left"
                                  else WD_ALIGN_PARAGRAPH.RIGHT)
            img_para.paragraph_format.space_before = Pt(0)
            img_para.paragraph_format.space_after  = Pt(0)
            run = img_para.add_run()
            run.add_picture(str(img_path), width=Inches(img_in))

            # Vertically centre the image cell
            tc   = img_cell._tc
            tcPr = tc.get_or_add_tcPr()
            for old in tcPr.findall(qn("w:vAlign")):
                tcPr.remove(old)
            vAlign = OxmlElement("w:vAlign")
            vAlign.set(qn("w:val"), "center")
            tcPr.append(vAlign)

        # ── Decide which header builder to use ────────────────────────────────
        hdr_cfg      = self.config["header"]
        hdr_line_cfg = self.config.get("header_line")
        img_rel      = hdr_cfg.get("image", "")

        if img_rel:
            # Resolve path relative to the document source directory
            # _src_dir is set by add_content; fall back to cwd
            src_dir = getattr(self, "_src_dir", Path("."))
            img_path = (src_dir / img_rel) if not Path(img_rel).is_absolute() else Path(img_rel)
            if img_path.exists():
                _build_header_with_image(section.header, hdr_cfg,
                                          hdr_line_cfg, img_path)
            else:
                _build_text(section.header, hdr_cfg, hdr_line_cfg, "bottom")
        else:
            _build_text(section.header, hdr_cfg, hdr_line_cfg, "bottom")

        _build_text(section.footer, self.config["footer"],
                    self.config.get("footer_line"), "top")

    # ── horizontal rule ───────────────────────────────────────────────────────

    def _emit_hrule(self):
        """Render --- as a page break."""
        para = self.doc.add_paragraph()
        run  = para.add_run()
        br   = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run._element.append(br)

    # ── block dispatcher ──────────────────────────────────────────────────────

    def _emit_block(self, node, src: Path):
        t = node.__class__.__name__
        
        # Reset alert flag for non-alert content to break the chain
        # Don't reset for BlankLine (empty lines between alerts should preserve the chain)
        if t not in ("BlockQuote", "Quote", "Alert", "BlankLine"):
            self._last_was_alert = False

        if   t == "Heading":       self._emit_heading(node, src)
        elif t == "Paragraph":     self._emit_paragraph(node, src)
        elif t in ("FencedCode", "CodeBlock"): self._emit_code(node)
        elif t == "List":          self._emit_list(node, src)
        elif t in ("BlockQuote", "Quote", "Alert"): self._emit_blockquote(node, src)
        elif t == "Table":         self._emit_table(node)
        elif t == "ThematicBreak": self._emit_hrule()
        elif t == "BlankLine":     pass

    # ── heading ───────────────────────────────────────────────────────────────

    def _emit_heading(self, node, src: Path):
        raw_text = self._plain_text(node)

        # Extract optional {.notoc} / {.nonumber} attribute block at end of text
        # e.g.  ## My Section {.notoc}   or   ### Appendix {.notoc .nonumber}
        notoc    = bool(re.search(r'\{[^}]*\.notoc[^}]*\}',    raw_text))
        nonumber = bool(re.search(r'\{[^}]*\.nonumber[^}]*\}', raw_text))
        # Strip the {…} block for the bookmark slug (visual text already stripped by _fill_inline)
        clean_text = re.sub(r'\s*\{[^}]+\}\s*$', '', raw_text).strip()

        level = min(node.level, 6)
        para  = self.doc.add_paragraph(style=f"Heading {level}")
        
        # Check if we're in appendix mode
        if self._appendix_mode:
            # Generate appendix label
            appendix_label = self._generate_appendix_label(level)
            
            # Add label as the first run
            label_run = para.add_run(f"{appendix_label}. ")
            label_run.bold = True
            # Use appropriate color based on level (matching regular heading hierarchy)
            from docx.shared import RGBColor
            if level == 1:
                label_run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)  # Dark blue
            elif level == 2:
                label_run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)  # Medium blue
            else:
                label_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)  # Grey for deeper levels
            
            # Now add the heading content
            self._fill_inline(para, node, src)
            
            # Suppress regular numbering for appendix headings
            if self._heading_num_id:
                _suppress_num_pr(para)
        else:
            self._fill_inline(para, node, src, strip_attrs=True)  # strips {.nonumber}/{.notoc}

        # Hide from TOC: set outlineLvl=9 (outside the TOC \o "1-3" range)
        if notoc:
            pPr = para._p.get_or_add_pPr()
            existing = pPr.find(qn("w:outlineLvl"))
            if existing is not None:
                pPr.remove(existing)
            ol = OxmlElement("w:outlineLvl")
            ol.set(qn("w:val"), "9")
            pPr.append(ol)

        # Numbered headings: style-level numPr (baked in _setup_numbering) handles
        # normal headings automatically.  For {.nonumber} we override with numId=0
        # to suppress the inherited style-level numbering on this paragraph only.
        if self._heading_num_id and nonumber and not self._appendix_mode:
            _suppress_num_pr(para)

        # Bookmark for cross-references
        slug    = _slugify(clean_text)
        bm_name = f"h-{slug}"
        self._bm_id += 1
        _add_bookmark(para, self._bm_id, bm_name)
        self._bm_map[f"#{slug}"] = bm_name
        self._bm_map[bm_name]    = bm_name

    def _generate_appendix_label(self, level: int) -> str:
        """Generate appendix label based on level and current counters.
        
        Level 1 (H1): Appendix A, Appendix B, Appendix C...
        Level 2 (H2): A1, A2, A3...
        Level 3 (H3): A1.1, A1.2, A1.3...
        Level 4+ (H4-H6): A1.1.1, A1.1.2, etc.
        
        Returns the label string (e.g., "Appendix A", "A1", "A1.1")
        """
        if level == 1:
            # H1: Appendix A, B, C...
            self._appendix_letter = chr(ord('A') + self._appendix_counters[0])
            self._appendix_counters[0] += 1
            # Reset sub-counters
            for i in range(1, 6):
                self._appendix_counters[i] = 0
            return f"Appendix {self._appendix_letter}"
        else:
            # H2-H6: Build hierarchical numbering
            # Increment counter for current level
            self._appendix_counters[level - 1] += 1
            # Reset deeper levels
            for i in range(level, 6):
                self._appendix_counters[i] = 0
            
            # Build label: A1, A1.1, A1.1.1, etc.
            label = self._appendix_letter
            for i in range(1, level):
                label += f".{self._appendix_counters[i]}"
            return label

    def _emit_cover_heading(self, node, src: Path):
        """Heading on the cover page — uses Cover Title / Subtitle / Body styles
        which have no outline_level, so the TOC field (\\o "1-3") never picks them up.
        No w:outlineLvl hack needed; the style itself is outside the heading family."""
        text  = self._plain_text(node)
        level = min(node.level, 6)

        # Map heading level → cover style (level 1 = title, 2 = subtitle, rest = body)
        cover_style = {1: "Cover Title", 2: "Cover Subtitle"}.get(level, "Cover Body")
        para = self.doc.add_paragraph(style=cover_style)
        self._fill_inline(para, node, src)
        
        # Apply styling overrides if set for this cover page block
        if self._current_alignment is not None:
            para.alignment = self._current_alignment
        if self._current_font_size is not None:
            for run in para.runs:
                run.font.size = Pt(self._current_font_size)
        if self._current_font_color is not None:
            for run in para.runs:
                run.font.color.rgb = self._current_font_color

        # No numPr — cover page headings are never numbered
        # No outlineLvl override needed — style has no outline_level at all

        # Register bookmark so internal links can target the cover title
        slug    = _slugify(text)
        bm_name = f"h-{slug}"
        self._bm_id += 1
        _add_bookmark(para, self._bm_id, bm_name)
        self._bm_map[f"#{slug}"] = bm_name
        self._bm_map[bm_name]    = bm_name

    # ── paragraph ─────────────────────────────────────────────────────────────

    def _emit_paragraph(self, node, src: Path):
        text = self._plain_text(node)

        # ── figure caption: *Figure: description {#anchor}*  (anchor optional)
        m = re.match(r"^\*?Figure:\s+(.*?)(?:\s+\{#([\w-]+)\})?\s*\*?$", text.strip())
        if m:
            self._emit_fig_caption(m.group(2), m.group(1).strip())
            return

        # ── table caption: *Table: description {#anchor}*  (anchor optional)
        m2 = re.match(r"^\*?Table:\s+(.*?)(?:\s+\{#([\w-]+)\})?\s*\*?$", text.strip())
        if m2:
            self._emit_tbl_caption(m2.group(2), m2.group(1).strip())
            return

        # check for image-only paragraph (may be followed by a RawText attr node)
        children = list(node.children) if hasattr(node, "children") and isinstance(node.children, list) else []
        if children and children[0].__class__.__name__ == "Image":
            # collect any trailing RawText siblings as attribute string
            attr_text = ""
            for c in children[1:]:
                if c.__class__.__name__ == "RawText":
                    attr_text += c.children if isinstance(c.children, str) else ""
            self._emit_image_para(children[0], src, attr_text)
            return

        # regular paragraph
        clean = re.sub(r"\s*\{[^}]+\}\s*$", "", text).strip()
        if not clean:
            return
        para = self.doc.add_paragraph()
        self._fill_inline(para, node, src)
        
        # Apply styling overrides if set for this cover page block
        if self._current_alignment is not None:
            para.alignment = self._current_alignment
        if self._current_font_size is not None:
            for run in para.runs:
                run.font.size = Pt(self._current_font_size)
        if self._current_font_color is not None:
            for run in para.runs:
                run.font.color.rgb = self._current_font_color

    def _emit_fig_caption(self, anchor: Optional[str], desc: str):
        """Render a figure caption.
        
        anchor: the name from {#anchor} (without #), or None if absent.
        desc:   the description text.
        
        Number is resolved from _label_map if anchor is present, otherwise
        the internal _fig_count is incremented.
        """
        if anchor and f"#{anchor}" in self._label_map:
            fig_num = self._label_map[f"#{anchor}"][1]
        else:
            self._fig_count += 1
            fig_num = self._fig_count

        bm_name = f"fig-{anchor}" if anchor else f"fig-{fig_num}"

        para = self.doc.add_paragraph(style="Caption")
        para.alignment = self._last_img_alignment

        # Don't set bold/italic explicitly — the Caption paragraph style
        # already defines both. Setting them on runs causes Word to strip
        # them on save (redundant explicit = style inheritance), which
        # produces false positives in the review diff.
        label = f"Figure {fig_num}:"
        para.add_run(label)
        if desc:
            para.add_run(" " + desc)

        self._bm_id += 1
        _add_bookmark(para, self._bm_id, bm_name)
        if anchor:
            self._bm_map[f"#{anchor}"] = bm_name

    def _emit_tbl_caption(self, anchor: Optional[str], desc: str):
        """Render a table caption.
        
        anchor: the name from {#anchor} (without #), or None if absent.
        desc:   the description text.
        """
        if anchor and f"#{anchor}" in self._label_map:
            tbl_num = self._label_map[f"#{anchor}"][1]
        else:
            self._tbl_count += 1
            tbl_num = self._tbl_count

        bm_name = f"tbl-{anchor}" if anchor else f"tbl-{tbl_num}"

        para = self.doc.add_paragraph(style="Caption")
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Don't set bold/italic explicitly — Caption style handles it.
        label = f"Table {tbl_num}:"
        para.add_run(label)
        if desc:
            para.add_run(" " + desc)

        self._bm_id += 1
        _add_bookmark(para, self._bm_id, bm_name)
        if anchor:
            self._bm_map[f"#{anchor}"] = bm_name

    # ── code block ────────────────────────────────────────────────────────────

    def _emit_code(self, node):
        # extract text
        code_text = ""
        children  = node.children if hasattr(node, "children") else []
        if isinstance(children, str):
            code_text = children
        elif isinstance(children, list) and children:
            c = children[0]
            if hasattr(c, "children") and isinstance(c.children, str):
                code_text = c.children
            else:
                code_text = str(c)

        para = self.doc.add_paragraph(style="Code")
        run  = para.add_run(code_text.rstrip("\n"))
        _cc = self.config.get("styles", {}).get("code", {})
        run.font.name = _cc.get("font_name", "Courier New")
        run.font.size = Pt(float(_cc.get("font_size_pt", 9)))

    # ── list ──────────────────────────────────────────────────────────────────

    def _emit_list(self, node, src: Path, level: int = 0,
                   forced_num_id: Optional[int] = None):
        """Emit a list node.

        For ordered and unordered lists, a fresh numId is allocated at the top
        level (level==0) and passed down to all nested levels so they share the
        same abstractNum but the counter is independent from any other list or
        from headings.  Explicit paragraph-level numPr is always written so it
        can never be accidentally inherited from a style that conflicts with
        heading numbering.
        """
        ordered = getattr(node, "ordered", False)

        # Determine the numId for list items at this nesting depth.
        if ordered:
            if forced_num_id is not None:
                # nested level — reuse parent's numId (same counter family)
                list_num_id = forced_num_id
            elif self._ol_abs_id is not None:
                # top-level ordered list — allocate a brand-new numId so this
                # list starts at 1 and shares nothing with any heading counter
                list_num_id = self._new_ol_num_id()
            else:
                list_num_id = None  # numbering disabled globally
        else:
            # unordered (bullet) list
            if forced_num_id is not None:
                # nested level — reuse parent's numId (same bullet family)
                list_num_id = forced_num_id
            elif self._ul_abs_id is not None:
                # top-level unordered list — allocate a brand-new numId
                list_num_id = self._new_ul_num_id()
            else:
                list_num_id = None  # numbering disabled globally

        for item in node.children:
            inline_nodes = []
            nested       = []
            for child in (item.children if hasattr(item, "children") and isinstance(item.children, list) else []):
                if child.__class__.__name__ == "List":
                    nested.append(child)
                elif child.__class__.__name__ == "Paragraph":
                    inline_nodes.append(child)
                else:
                    inline_nodes.append(child)

            style = ("List Number" if ordered else "List Bullet") + (f" {level+1}" if level > 0 else "")
            try:
                para = self.doc.add_paragraph(style=style)
            except Exception:
                para = self.doc.add_paragraph(style="List Bullet" if not ordered else "List Number")

            para.clear()
            for child in inline_nodes:
                self._fill_inline(para, child, src)

            # Attach explicit numPr for both ordered and unordered lists.
            # This ensures the paragraph uses our dedicated abstractNum counter
            # rather than anything style-linked (which might share state with
            # the heading abstractNum counter).
            if list_num_id is not None:
                _add_num_pr(para, list_num_id, level)

            for nested_list in nested:
                self._emit_list(
                    nested_list, src, level + 1,
                    forced_num_id=list_num_id,
                )

    # ── blockquote ────────────────────────────────────────────────────────────

    def _emit_blockquote(self, node, src: Path):
        """Render blockquote with optional GitHub-style alert types.
        
        Supports:
        - [!CAUTION] or [!DANGER] → Red border
        - [!WARNING] → Orange border  
        - [!NOTE] → Blue border (default blockquote)
        - [!TIP] or [!SUCCESS] → Green border
        
        For alerts, uses a 1x1 table with colored left border for clean visual separation.
        """
        from docx.shared import RGBColor
        
        # Define alert type colors (border color)
        _ac = self.config.get("styles", {}).get("alerts", {})
        ALERT_STYLES = {
            'caution':  _ac.get("caution_color", "DC3545").lstrip("#"),
            'danger':   _ac.get("caution_color", "DC3545").lstrip("#"),
            'warning':  _ac.get("warning_color", "FFA500").lstrip("#"),
            'note':     _ac.get("note_color",    "2E75B6").lstrip("#"),
            'tip':      _ac.get("tip_color",     "28A745").lstrip("#"),
            'success':  _ac.get("tip_color",     "28A745").lstrip("#"),
        }
        
        # Check if this is an Alert node (GitHub-style [!TYPE] syntax)
        alert_type = None
        border_color = '2E75B6'  # Default blue
        
        node_type = node.__class__.__name__
        if node_type == 'Alert' and hasattr(node, 'alert_type'):
            alert_type = node.alert_type.lower()
            if alert_type in ALERT_STYLES:
                border_color = ALERT_STYLES[alert_type]
        
        is_alert_node = node_type == 'Alert' and alert_type is not None
        
        if is_alert_node:
            # For alerts: Use a 1x1 table with colored left border
            # This creates clean visual boxes that don't "melt" together
            table = self.doc.add_table(rows=1, cols=1)
            table.autofit = False
            # Indent table to match code/quote left inset (0.15in = 216 twips)
            _tblPr = table._tbl.tblPr
            for _old in _tblPr.findall(qn("w:tblInd")):
                _tblPr.remove(_old)
            _ind = OxmlElement("w:tblInd")
            _ind.set(qn("w:w"), "216")
            _ind.set(qn("w:type"), "dxa")
            _tblPr.append(_ind)
            
            # Get the cell — match width of code blocks (content width minus 0.15in each side)
            cell = table.cell(0, 0)
            content_in = self._img.inches(self._img.content_width_emu)
            alert_in   = content_in - 0.30   # same inset as code/quote styles
            cell.width = Inches(alert_in)
            
            # Set cell margins
            cell_margin = 100  # ~1.4mm
            tcPr = cell._tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for side in ['top', 'left', 'bottom', 'right']:
                margin = OxmlElement(f'w:{side}')
                margin.set(qn('w:w'), str(cell_margin))
                margin.set(qn('w:type'), 'dxa')
                tcMar.append(margin)
            tcPr.append(tcMar)
            
            # Add colored left border only + light grey background
            tcBorders = OxmlElement('w:tcBorders')
            
            # Left border - colored and visible
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '24')  # Thick border for visibility
            left.set(qn('w:color'), border_color)
            tcBorders.append(left)
            
            # No other borders
            for side in ['top', 'right', 'bottom']:
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'), 'nil')
                tcBorders.append(b)
            
            tcPr.append(tcBorders)
            
            # Add light grey background shading
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            _alert_bg = self.config.get("styles", {}).get("alerts", {}).get("background", "F5F5F5").lstrip("#")
            shd.set(qn('w:fill'), _alert_bg)
            tcPr.append(shd)
            
            # Add content to cell
            cell_para = cell.paragraphs[0]
            cell_para.paragraph_format.space_before = Pt(6)
            cell_para.paragraph_format.space_after = Pt(6)
            
            # Add label and content
            children = node.children if hasattr(node, 'children') and isinstance(node.children, list) else []
            for i, child in enumerate(children):
                if child.__class__.__name__ == 'Paragraph':
                    if i == 0 and alert_type:
                        # First paragraph: Add label in bold with colored text
                        label_run = cell_para.add_run(f'[{alert_type.upper()}] ')
                        label_run.bold = True
                        label_run.font.color.rgb = RGBColor(
                            int(border_color[0:2], 16),
                            int(border_color[2:4], 16),
                            int(border_color[4:6], 16)
                        )
                    self._fill_inline(cell_para, child, src)
                    if i < len(children) - 1:
                        # Add line break between paragraphs within the same cell
                        cell_para.add_run().add_break()
            
            # Add spacing after alert table for visual separation between consecutive alerts
            # 0pt gap - no extra spacing, relies on grey backgrounds for separation
            spacer = self.doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(0)  # No gap - boxes touch but grey bg provides separation
        else:
            # Regular blockquote: Use paragraph style
            children = node.children if hasattr(node, 'children') and isinstance(node.children, list) else []
            for child in children:
                if child.__class__.__name__ == 'Paragraph':
                    para = self.doc.add_paragraph(style='Block Quote')
                    self._fill_inline(para, child, src)
                else:
                    self._emit_block(child, src)

    # ── table ─────────────────────────────────────────────────────────────────

    def _emit_table(self, node):
        # GFM: node.head = header row, node.children = all rows
        if not hasattr(node, "head") or not node.head:
            return

        # Keep both plain-text (for merge detection) and node grids (for rendering)
        header_nodes = list(node.head.children)
        body_row_nodes = [list(row.children) for row in list(node.children)[1:]]

        header_cells = [self._plain_text(c) for c in header_nodes]
        body_rows    = [[self._plain_text(c) for c in row] for row in body_row_nodes]

        has_header = any(cell.strip() for cell in header_cells)

        self._tbl_count += 1

        # ── Resolve merge plan ────────────────────────────────────────────────
        clean_grid, merge_grid = build_merge_plan(header_cells, body_rows, has_header)

        nrows = len(clean_grid)
        ncols = len(clean_grid[0]) if clean_grid else len(header_cells)

        table = self.doc.add_table(rows=nrows, cols=ncols)

        # ── Fill cell content via _fill_inline to preserve formatting ─────────
        # Build a flat list of all cell nodes in document order for rendering.
        # Fill cell content, skipping cells consumed by merges.
        all_node_rows = ([header_nodes] if has_header else []) + body_row_nodes
        for ri, node_row in enumerate(all_node_rows):
            for ci, cell_node in enumerate(node_row):
                if ci >= ncols:
                    continue
                # Skip cells consumed by a merge: their anchor points elsewhere
                anc_r, anc_c = merge_grid[ri][ci][0], merge_grid[ri][ci][1]
                if (anc_r, anc_c) != (ri, ci):
                    continue   # consumed cell — leave blank, merge will cover it
                doc_cell = table.rows[ri].cells[ci]
                for p in doc_cell.paragraphs:
                    p.clear()
                para = doc_cell.paragraphs[0]
                self._fill_inline(para, cell_node, self._src_dir)

        # Apply merges and per-cell alignment
        apply_merges(table, merge_grid, nrows, ncols)

        # Style the table
        raw_rows = ([header_cells] if has_header else []) + body_rows
        if is_image_table(raw_rows):
            format_image_table(table)
        else:
            format_data_table(table, has_header=has_header,
                              table_cfg=self.config.get("styles"))

        self._last_table = table

    # ── image group (:::figures) ──────────────────────────────────────────────

    def _emit_image_group(self, block_text: str, src: Path):
        """Render a :::figures block as a side-by-side row of images.

        Each image gets a sub-caption  a) / b) / c) …  inside its cell.
        Place a normal *Figure N: …* caption below the block for the group label.

        Syntax:
            :::figures
            ![First description](img1.png){.medium}
            ![Second description](img2.png){.medium}
            :::
            *Figure 3: Overall caption.*
        """
        lines = block_text.strip().splitlines()
        image_lines = [l.strip() for l in lines if l.strip().startswith("![")]

        if not image_lines:
            return

        n      = len(image_lines)
        labels = [chr(ord("a") + i) for i in range(n)]

        # Total text-area width in dxa (1 dxa = 635 EMU).  Using 5000 pct table
        # is more reliable but we also set explicit dxa widths for cell sizing.
        TOTAL_DXA    = 8748          # ≈ full A4 content width at 2.54 cm margins
        dxa_per_col  = TOTAL_DXA // n

        # Each cell has 60 dxa left + 60 dxa right = 120 dxa of horizontal padding
        # (set by format_image_table).  Images must be sized to the usable inner
        # width or they overflow into the next column / onto the next page.
        CELL_MARGIN_H_DXA = 120          # 60 left + 60 right
        EMU_PER_DXA       = 635          # 1 dxa = 635 EMU (914400 EMU/in ÷ 1440 dxa/in)
        col_width_emu = (self._img.content_width_emu // n) - CELL_MARGIN_H_DXA * EMU_PER_DXA

        # Borderless table, one row, n columns
        table = self.doc.add_table(rows=1, cols=n)
        format_image_table(table)

        # Set table to 100 % of content width so columns fill the page evenly
        tbl  = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"),    "5000")
        tblW.set(qn("w:type"), "pct")
        existing_tblW = tblPr.find(qn("w:tblW"))
        if existing_tblW is not None:
            tblPr.remove(existing_tblW)
        tblPr.append(tblW)

        # Set equal column widths
        for cell in table.rows[0].cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW  = OxmlElement("w:tcW")
            tcW.set(qn("w:w"),    str(dxa_per_col))
            tcW.set(qn("w:type"), "dxa")
            existing = tcPr.find(qn("w:tcW"))
            if existing is not None:
                tcPr.remove(existing)
            tcPr.append(tcW)

        for i, line in enumerate(image_lines):
            m = re.match(r'!\[([^\]]*)\]\(([^)]*)\)(.*)', line)
            if not m:
                continue
            alt_text    = m.group(1).strip()
            path_str    = m.group(2).strip()
            extra_attrs = m.group(3).strip()
            combined    = alt_text + " " + extra_attrs

            # Parse size class / width attribute
            size_class = None
            width_attr = None
            m_cls = re.search(r'\{[^}]*\.(\w+)[^}]*\}', combined)
            if m_cls and m_cls.group(1) in self._img.size_classes:
                size_class = m_cls.group(1)
            m_w = re.search(r'\{width=([^}]+)\}', combined)
            if m_w:
                width_attr = f"width={m_w.group(1)}"

            cell     = table.rows[0].cells[i]
            img_para = cell.paragraphs[0]
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            result = self._img.load(path_str, src.parent if src.is_file() else src)
            if result:
                fp, w_px, h_px = result
                w_emu, h_emu   = self._img.calc_emu(w_px, h_px, size_class, width_attr)
                # Cap to column width so images never overflow
                if w_emu > col_width_emu:
                    aspect = h_px / w_px if w_px else 1.0
                    w_emu  = col_width_emu
                    h_emu  = int(w_emu * aspect)
                run = img_para.add_run()
                try:
                    run.add_picture(str(fp), width=Inches(self._img.inches(w_emu)))
                except Exception:
                    img_para.add_run(f"[{path_str}]")
            else:
                img_para.add_run(f"[Image not found: {path_str}]")

            # Sub-caption:  a) Description text
            sub_para = cell.add_paragraph()
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = sub_para.add_run(f"{labels[i]}) {alt_text}")
            r.italic          = True
            r.font.size       = Pt(9)
            r.font.color.rgb  = RGBColor(0x55, 0x55, 0x55)

        # Keep the last paragraph in every cell with the next block (the figure
        # caption), so Word never page-breaks between the image group and its label.
        for cell in table.rows[-1].cells:
            for p in cell.paragraphs:
                p.paragraph_format.keep_with_next = True

        # The Figure N: caption that follows will inherit CENTER alignment
        self._last_img_alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _fill_inline_fmt(self, para, node, src, bold=False, italic=False, strip_attrs=False):
        """Recurse into an inline formatting node accumulating bold/italic state.

        Handles nested cases like ***bold+italic*** which marko parses as
        Emphasis(StrongEmphasis(RawText)) — the outer node is Emphasis (italic)
        but we need to carry bold=True down from the StrongEmphasis inside.
        """
        children = node.children if hasattr(node, "children") else []
        if isinstance(children, str):
            if children:
                r = para.add_run(children)
                r.bold = bold; r.italic = italic
            return
        if not isinstance(children, list):
            return
        for child in children:
            ct = child.__class__.__name__
            if ct == "RawText":
                raw = child.children if isinstance(child.children, str) else ""
                if strip_attrs:
                    raw = re.sub(r"\s*\{[^}]+\}\s*$", "", raw)
                if raw:
                    r = para.add_run(raw)
                    r.bold = bold; r.italic = italic
            elif ct in ("Emphasis", "Em"):
                self._fill_inline_fmt(para, child, src, bold=bold, italic=True,
                                      strip_attrs=strip_attrs)
            elif ct in ("Strong", "StrongEmphasis"):
                self._fill_inline_fmt(para, child, src, bold=True, italic=italic,
                                      strip_attrs=strip_attrs)
            elif ct == "Strikethrough":
                txt = self._plain_text(child)
                if txt:
                    r = para.add_run(txt)
                    r.bold = bold; r.italic = italic; r.font.strike = True
            elif ct in ("InlineCode", "Code", "CodeSpan"):
                txt = self._plain_text(child)
                if txt:
                    r = para.add_run(txt)
                    r.bold = bold; r.italic = italic
                    r.font.name = "Courier New"
            else:
                txt = self._plain_text(child)
                if txt:
                    r = para.add_run(txt)
                    r.bold = bold; r.italic = italic

        # ── image paragraph ───────────────────────────────────────────────────────

    def _emit_image_para(self, img_node, src: Path, extra_attrs: str = ""):
        """Emit a paragraph containing one image, with optional alignment and size."""
        path_str  = img_node.dest if hasattr(img_node, "dest") else ""
        alt_text  = self._plain_text(img_node)

        # combine alt text + any trailing attribute tokens
        combined = alt_text + " " + extra_attrs

        # parse size class: {.xs}, {.small}, {.medium}, {.large}, {.xl}
        size_class = None
        width_attr = None
        # Extract all classes from braces, find first that matches size classes
        classes = re.findall(r"\{\.([^.}\s]+)", combined)
        size_class = next((c for c in classes if c in self._img.size_classes), None)
        m_w = re.search(r"\{width=([^}]+)\}", combined)
        if m_w:
            width_attr = f"width={m_w.group(1)}"

        # parse alignment: {.left}, {.center}, {.right}
        alignment = WD_ALIGN_PARAGRAPH.CENTER  # default
        if re.search(r"\{[^}]*\.left[^}]*\}", combined):
            alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif re.search(r"\{[^}]*\.right[^}]*\}", combined):
            alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif re.search(r"\{[^}]*\.center[^}]*\}", combined):
            alignment = WD_ALIGN_PARAGRAPH.CENTER

        result = self._img.load(path_str, src.parent if src.is_file() else src)
        if not result:
            self.doc.add_paragraph(f"[Image not found: {path_str}]")
            return
        full_path, w_px, h_px = result
        w_emu, h_emu = self._img.calc_emu(w_px, h_px, size_class, width_attr)

        self._last_img_alignment = alignment  # captions inherit this

        para = self.doc.add_paragraph()
        para.alignment = alignment
        # Keep this paragraph with the next one (the caption) — prevents Word
        # from inserting a page break between the image and its caption.
        para.paragraph_format.keep_with_next = True
        run = para.add_run()
        try:
            run.add_picture(str(full_path), width=Inches(self._img.inches(w_emu)))
        except Exception:
            para.add_run(f"[Could not load: {path_str}]")

    # ── inline content ────────────────────────────────────────────────────────

    def _fill_inline(self, para, node, src: Path, children_override=None, strip_attrs=False):
        """Recursively fill a paragraph with inline-formatted runs."""
        children = children_override
        if children is None:
            if hasattr(node, "children"):
                children = node.children
            else:
                return

        if isinstance(children, str):
            para.add_run(children)
            return

        if not isinstance(children, list):
            return

        for i, child in enumerate(children):
            ct = child.__class__.__name__

            if ct == "RawText":
                raw = child.children if isinstance(child.children, str) else str(child)
                # Skip standalone attribute tokens:
                # - CSS class blocks: {.medium} {.large .center}
                # - Table cell merge/align attrs: {cs=2} {rs=3} {ha=c} {va=t}
                if re.fullmatch(r"\s*\{[^}]+\}\s*", raw):
                    continue
                # Also strip trailing cell attributes from the last text run
                # e.g. "Region {cs=2}" → strip the " {cs=2}" suffix
                raw = re.sub(r"\s*\{(?:cs|rs|ha|va)=[^}]+\}\s*$", "", raw)
                # In heading context, also strip trailing attribute block e.g. {.nonumber}
                if strip_attrs:
                    raw = re.sub(r"\s*\{[^}]+\}\s*$", "", raw)
                if raw:
                    para.add_run(raw)

            elif ct in ("Emphasis", "Em"):
                # Recurse with italic=True so nested StrongEmphasis (***text***)
                # correctly carries bold down through the Emphasis wrapper
                self._fill_inline_fmt(para, child, src, italic=True,
                                      strip_attrs=strip_attrs)

            elif ct == "Strong":
                self._fill_inline_fmt(para, child, src, bold=True,
                                      strip_attrs=strip_attrs)

            elif ct == "StrongEmphasis":
                # StrongEmphasis = **bold** (bold only, not italic)
                # ***bold italic*** parses as Emphasis(StrongEmphasis(...))
                # which is handled by _fill_inline_fmt with italic=True carried down
                txt = self._plain_text(child)
                r   = para.add_run(txt); r.bold = True

            elif ct == "Strikethrough":
                txt = self._plain_text(child)
                r   = para.add_run(txt); r.font.strike = True

            elif ct in ("InlineCode", "Code", "CodeSpan"):
                txt = self._plain_text(child)
                r   = para.add_run(txt)
                r.font.name = "Courier New"; r.font.size = Pt(10)

            elif ct == "Link":
                url  = child.dest if hasattr(child, "dest") else "#"
                txt  = self._plain_text(child)
                if url.startswith("http"):
                    self._add_ext_hyperlink(para, txt, url)
                else:
                    # internal citation: resolve bookmark
                    # If the URL is a named figure/table anchor, auto-resolve
                    # the display text to include the number e.g. "Figure 1"
                    if url in self._label_map:
                        kind, num = self._label_map[url]
                        # Replace generic "Figure" or "Table" in txt with numbered version
                        resolved_txt = re.sub(
                            r'\b(Figure|Table)\b',
                            f"{kind} {num}",
                            txt,
                            count=1
                        )
                        anchor = url[1:]  # strip leading #
                        bm = f"{'fig' if kind == 'Figure' else 'tbl'}-{anchor}"
                        self._add_int_hyperlink(para, resolved_txt, bm)
                    elif url in self._bm_map:
                        bm = self._bm_map[url]
                        self._add_int_hyperlink(para, txt, bm)
                    elif url.startswith("#"):
                        slug = url[1:]
                        bm = f"h-{slug}"
                        self._add_int_hyperlink(para, txt, bm)
                    else:
                        self._add_int_hyperlink(para, txt, url)

            elif ct == "Image":
                # inline image inside a mixed paragraph
                path_str = child.dest if hasattr(child, "dest") else ""
                size_class = None; width_attr = None
                # look at next sibling for attribute token
                if i + 1 < len(children):
                    nxt = children[i + 1]
                    if nxt.__class__.__name__ == "RawText":
                        raw_nxt = nxt.children if isinstance(nxt.children, str) else ""
                        m_cls = re.search(r"\{\.(\w+)\}", raw_nxt)
                        if m_cls and m_cls.group(1) in self._img.size_classes:
                            size_class = m_cls.group(1)
                        m_w = re.search(r"\{width=([^}]+)\}", raw_nxt)
                        if m_w:
                            width_attr = f"width={m_w.group(1)}"
                result = self._img.load(path_str, src.parent if src.is_file() else src)
                if result:
                    fp, w_px, h_px = result
                    w_emu, h_emu   = self._img.calc_emu(w_px, h_px, size_class, width_attr)
                    try:
                        run = para.add_run()
                        run.add_picture(str(fp), width=Inches(self._img.inches(w_emu)))
                    except Exception:
                        para.add_run(f"[{path_str}]")
                else:
                    para.add_run(f"[Image: {path_str}]")

            elif ct in ("LineBreak", "HardBreak"):
                run = para.add_run()
                br  = OxmlElement("w:br")
                run._element.append(br)

            elif ct in ("SoftBreak", "SoftLineBreak"):
                para.add_run(" ")

            elif hasattr(child, "children"):
                self._fill_inline(para, child, src)

    # ── hyperlink helpers ─────────────────────────────────────────────────────

    def _add_ext_hyperlink(self, para, text: str, url: str):
        """Clickable external hyperlink: run lives inside w:hyperlink."""
        r_id = para.part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hl = OxmlElement("w:hyperlink")
        hl.set(qn("r:id"), r_id)
        r  = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        for tag, attrib in [
            ("w:rStyle", {qn("w:val"): "Hyperlink"}),
            ("w:color",  {qn("w:val"): "0563C1"}),
            ("w:u",      {qn("w:val"): "single"}),
        ]:
            el = OxmlElement(tag)
            for k, v in attrib.items(): el.set(k, v)
            rPr.append(el)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve"); t.text = text
        r.append(t)
        hl.append(r)
        para._p.append(hl)

    def _add_int_hyperlink(self, para, text: str, anchor: str):
        """Clickable internal cross-reference (jumps to bookmark)."""
        hl = OxmlElement("w:hyperlink")
        hl.set(qn("w:anchor"), anchor)
        r  = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        for tag, attrib in [
            ("w:rStyle", {qn("w:val"): "Hyperlink"}),
            ("w:color",  {qn("w:val"): "0563C1"}),
            ("w:u",      {qn("w:val"): "single"}),
        ]:
            el = OxmlElement(tag)
            for k, v in attrib.items(): el.set(k, v)
            rPr.append(el)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve"); t.text = text
        r.append(t)
        hl.append(r)
        para._p.append(hl)

    # ── text extraction ───────────────────────────────────────────────────────

    def _plain_text(self, node) -> str:
        parts = []
        def collect(n):
            if hasattr(n, "children"):
                if isinstance(n.children, str):
                    parts.append(n.children)
                elif isinstance(n.children, list):
                    for c in n.children: collect(c)
        collect(node)
        return "".join(parts)
